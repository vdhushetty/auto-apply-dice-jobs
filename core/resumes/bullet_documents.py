"""Conservative, layout-preserving DOCX bullet rewriting.

The model-facing planner works with small immutable ``EditableBullet`` values. This
module owns the corresponding Word paragraphs and treats the planner output as
untrusted: only known bullet IDs may change, at most one clone may be inserted for
an edited paragraph, and every other package part must remain structurally equal
(and binary parts byte-for-byte equal).
"""

from __future__ import annotations

import os
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from docx.text.paragraph import Paragraph
from lxml import etree  # type: ignore[import-untyped]

from .bullet_curator import (
    MAX_EDITED_BULLETS,
    MAX_NET_NEW_BULLETS,
    MAX_REPLACEMENT_BULLET_CHARS,
    EditableBullet,
    ValidatedBulletRewritePlan,
)
from .documents import validate_resume_path
from .models import ResumeTailoringError

WORDPROCESSINGML_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DOCUMENT_PART = "word/document.xml"

_W_P = f"{{{WORDPROCESSINGML_NAMESPACE}}}p"
_W_P_PR = f"{{{WORDPROCESSINGML_NAMESPACE}}}pPr"
_W_R = f"{{{WORDPROCESSINGML_NAMESPACE}}}r"
_W_R_PR = f"{{{WORDPROCESSINGML_NAMESPACE}}}rPr"
_W_T = f"{{{WORDPROCESSINGML_NAMESPACE}}}t"

_ALLOWED_PARAGRAPH_CHILDREN = frozenset(
    {
        _W_P_PR,
        _W_R,
        f"{{{WORDPROCESSINGML_NAMESPACE}}}bookmarkStart",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}bookmarkEnd",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}proofErr",
    }
)
_ALLOWED_RUN_CHILDREN = frozenset({_W_R_PR, _W_T})
_UNSUPPORTED_BULLET_TAGS = frozenset(
    {
        f"{{{WORDPROCESSINGML_NAMESPACE}}}altChunk",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}br",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}commentRangeStart",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}commentRangeEnd",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}drawing",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}fldChar",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}fldSimple",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}hyperlink",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}ins",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}del",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}instrText",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}object",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}pict",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}sdt",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}smartTag",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}tab",
    }
)

_ACTION_VERBS = frozenset(
    {
        "achieved",
        "administered",
        "analyzed",
        "architected",
        "automated",
        "built",
        "collaborated",
        "configured",
        "coordinated",
        "created",
        "delivered",
        "deployed",
        "designed",
        "developed",
        "engineered",
        "enhanced",
        "established",
        "executed",
        "generated",
        "implemented",
        "improved",
        "integrated",
        "launched",
        "led",
        "maintained",
        "managed",
        "mentored",
        "migrated",
        "modernized",
        "monitored",
        "optimized",
        "orchestrated",
        "owned",
        "partnered",
        "planned",
        "processed",
        "programmed",
        "rebuilt",
        "reduced",
        "refactored",
        "resolved",
        "scaled",
        "streamlined",
        "supported",
        "transformed",
        "validated",
        "wrote",
    }
)
_EXPERIENCE_SECTION_MARKERS = frozenset(
    {
        "career highlights",
        "career history",
        "employment history",
        "experience",
        "professional experience",
        "relevant experience",
        "work experience",
        "work history",
    }
)
_PROJECT_SECTION_MARKERS = frozenset(
    {
        "project experience",
        "projects",
        "selected projects",
    }
)
_SECTION_CATEGORIES: dict[str, frozenset[str]] = {
    "certifications": frozenset(
        {
            "certifications",
            "certifications and licenses",
            "certifications licenses",
            "licenses",
        }
    ),
    "education": frozenset(
        {
            "education",
            "education and training",
            "education training",
            "training",
        }
    ),
    "other": frozenset(
        {
            "awards",
            "awards and honors",
            "awards honors",
            "publications",
        }
    ),
    "skills": frozenset(
        {
            "skills",
            "skills and tools",
            "skills tools",
            "technical skills",
            "technical skills and tools",
            "technical skills tools",
            "technologies",
        }
    ),
    "summary": frozenset(
        {
            "career objective",
            "professional profile",
            "professional summary",
            "qualifications",
            "summary",
        }
    ),
}
_LEADING_MANUAL_MARKER = re.compile(r"^[\u2022\u25cf\u25aa\u25e6*\-\u2013\u2014]\s+")
_REPLACEMENT_MARKER = re.compile(r"^(?:[\u2022\u25cf\u25aa\u25e6*\-\u2013\u2014]|\d+[.)])\s+")


@dataclass(frozen=True)
class _BulletRecord:
    editable: EditableBullet
    paragraph: Paragraph
    paragraph_index: int


@dataclass(frozen=True)
class _ResumeItemRecord:
    """One structure-stable paragraph exposed to whole-resume optimization."""

    editable: EditableBullet
    paragraph: Paragraph
    paragraph_index: int
    preserved_prefix: str = ""


def _normalize_heading(text: str) -> str:
    return " ".join(re.sub(r"[^a-z0-9 ]+", " ", text.casefold()).split())


def _section_kind(text: str) -> str | None:
    normalized = _normalize_heading(text)
    if normalized in _PROJECT_SECTION_MARKERS:
        return "projects"
    if normalized in _EXPERIENCE_SECTION_MARKERS:
        return "experience"
    for category, markers in _SECTION_CATEGORIES.items():
        if normalized in markers:
            return category
    return None


def _style_name(paragraph: Paragraph) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = _style_name(paragraph).casefold().strip()
    match = re.match(r"heading\s+(\d+)$", style_name)
    if match:
        return int(match.group(1))
    if style_name in {"heading", "title", "subtitle"}:
        return 1
    return None


def _is_heading_like(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 100:
        return False
    if _heading_level(paragraph) is not None or _section_kind(text) is not None:
        return True
    return text.isupper() and len(text.split()) <= 10 and not text.endswith((".", ";"))


def _iter_body_paragraphs(document: Any) -> tuple[Paragraph, ...]:
    """Return body and table-cell paragraphs in their exact XML document order."""

    return tuple(Paragraph(element, document) for element in document.element.body.iter(_W_P))


def _has_numbering(paragraph: Paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None and paragraph_properties.numPr is not None:
        return True

    style = paragraph.style
    seen: set[str] = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        style_name = style.name.casefold()
        if "bullet" in style_name or re.search(r"\blist\b", style_name):
            return True
        style_properties = style.element.pPr
        if style_properties is not None and style_properties.numPr is not None:
            return True
        style = style.base_style
    return False


def _starts_with_action_verb(text: str) -> bool:
    if _LEADING_MANUAL_MARKER.match(text.lstrip()):
        return False
    match = re.match(r"([A-Za-z]+)", text.lstrip())
    return bool(match and match.group(1).casefold() in _ACTION_VERBS)


def _has_conservative_fallback_bullet_geometry(paragraph: Paragraph) -> bool:
    """Require direct bullet-like indentation for an unnumbered fallback paragraph."""

    left_indent = paragraph.paragraph_format.left_indent
    first_line_indent = paragraph.paragraph_format.first_line_indent
    return bool(
        (left_indent is not None and left_indent > 0)
        or (first_line_indent is not None and first_line_indent < 0)
    )


def _is_candidate_bullet(paragraph: Paragraph, *, in_experience_section: bool) -> bool:
    text = paragraph.text.strip()
    if not text or _is_heading_like(paragraph):
        return False
    if not in_experience_section:
        return False
    if _has_numbering(paragraph):
        return True
    word_count = len(text.split())
    return (
        "|" not in text
        and "\t" not in text
        and text.endswith((".", ";"))
        and 4 <= word_count <= 120
        and len(text) <= MAX_REPLACEMENT_BULLET_CHARS
        and _has_conservative_fallback_bullet_geometry(paragraph)
        and _starts_with_action_verb(text)
    )


def _run_format_signature(run: Any) -> bytes:
    properties = run._r.rPr
    if properties is None:
        return b""
    return etree.tostring(properties, method="c14n")


def _validate_simple_bullet_paragraph(paragraph: Paragraph, bullet_id: str) -> None:
    descendants = {element.tag for element in paragraph._p.iter()}
    if descendants.intersection(_UNSUPPORTED_BULLET_TAGS):
        raise ResumeTailoringError(
            f"Bullet {bullet_id} contains a hyperlink, field, drawing, or other unsupported content."
        )
    if any(child.tag not in _ALLOWED_PARAGRAPH_CHILDREN for child in paragraph._p):
        raise ResumeTailoringError(f"Bullet {bullet_id} uses unsupported complex Word markup.")

    direct_runs = [child for child in paragraph._p if child.tag == _W_R]
    for run_element in direct_runs:
        if any(child.tag not in _ALLOWED_RUN_CHILDREN for child in run_element):
            raise ResumeTailoringError(f"Bullet {bullet_id} uses unsupported complex run content.")

    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        raise ResumeTailoringError(f"Bullet {bullet_id} has no safely editable text run.")
    if len({_run_format_signature(run) for run in text_runs}) > 1:
        raise ResumeTailoringError(
            f"Bullet {bullet_id} mixes text formatting and cannot be rewritten safely."
        )


def _mixed_skill_prefix(paragraph: Paragraph, item_id: str) -> tuple[str, str]:
    """Return a run-boundary label prefix and editable skill text without flattening styles."""

    descendants = {element.tag for element in paragraph._p.iter()}
    if descendants.intersection(_UNSUPPORTED_BULLET_TAGS):
        raise ResumeTailoringError(f"Skill item {item_id} contains unsupported Word content.")
    if any(child.tag not in _ALLOWED_PARAGRAPH_CHILDREN for child in paragraph._p):
        raise ResumeTailoringError(f"Skill item {item_id} uses unsupported complex Word markup.")
    direct_runs = [child for child in paragraph._p if child.tag == _W_R]
    if any(
        child.tag not in _ALLOWED_RUN_CHILDREN
        for run_element in direct_runs
        for child in run_element
    ):
        raise ResumeTailoringError(f"Skill item {item_id} uses unsupported complex run content.")

    text_runs = [run for run in paragraph.runs if run.text]
    leading = ""
    boundary_index = -1
    for index, run in enumerate(text_runs[:-1]):
        leading += run.text
        if leading.rstrip().endswith(":"):
            boundary_index = index
            break
    if boundary_index < 0 or len(leading.strip()) > 80:
        raise ResumeTailoringError(f"Skill item {item_id} has no safe label/content run boundary.")
    remaining = "".join(run.text for run in text_runs[boundary_index + 1 :])
    editable_text = remaining.strip()
    if not editable_text:
        raise ResumeTailoringError(f"Skill item {item_id} has no editable content after its label.")
    leading_whitespace = remaining[: len(remaining) - len(remaining.lstrip())]
    return f"{leading}{leading_whitespace}", editable_text


def _collect_bullet_records(document: Any) -> tuple[_BulletRecord, ...]:
    records: list[_BulletRecord] = []
    section_label = "other"
    section_kind = "other"
    group_index = 0
    previous_was_bullet = False

    for paragraph_index, paragraph in enumerate(_iter_body_paragraphs(document)):
        text = paragraph.text.strip()
        if not text:
            previous_was_bullet = False
            continue

        if _is_heading_like(paragraph):
            detected_kind = _section_kind(text)
            level = _heading_level(paragraph)
            if detected_kind is not None or level == 1:
                section_label = detected_kind or "other"
                section_kind = detected_kind or "other"
            previous_was_bullet = False
            continue

        if not _is_candidate_bullet(
            paragraph,
            in_experience_section=section_kind in {"experience", "projects"},
        ):
            previous_was_bullet = False
            continue

        bullet_id = f"bullet-{paragraph_index:04d}"
        try:
            _validate_simple_bullet_paragraph(paragraph, bullet_id)
        except ResumeTailoringError:
            # A complex paragraph is not exposed to the planner. It also breaks the
            # evidence group so content on either side cannot be combined through it.
            previous_was_bullet = False
            continue
        if not previous_was_bullet:
            group_index += 1
        group_id = f"group-{group_index:04d}"
        editable = EditableBullet(
            bullet_id=bullet_id,
            text=paragraph.text.strip(),
            section=section_label,
            group_id=group_id,
        )
        records.append(
            _BulletRecord(
                editable=editable,
                paragraph=paragraph,
                paragraph_index=paragraph_index,
            )
        )
        previous_was_bullet = True

    if not records:
        raise ResumeTailoringError("The resume did not contain any safely editable bullet points.")
    return tuple(records)


def collect_editable_bullets(document: Any) -> tuple[EditableBullet, ...]:
    """Collect simple resume bullets and assign deterministic paragraph-based IDs."""

    return tuple(record.editable for record in _collect_bullet_records(document))


def _collect_resume_item_records(document: Any) -> tuple[_ResumeItemRecord, ...]:
    """Collect safe summary, skills, experience, and project paragraphs in document order."""

    try:
        bullet_records = _collect_bullet_records(document)
    except ResumeTailoringError as exc:
        if "did not contain any safely editable bullet points" not in str(exc):
            raise
        bullet_records = ()
    bullets_by_index = {record.paragraph_index: record for record in bullet_records}
    records: list[_ResumeItemRecord] = []
    section_kind = "other"

    for paragraph_index, paragraph in enumerate(_iter_body_paragraphs(document)):
        bullet_record = bullets_by_index.get(paragraph_index)
        if bullet_record is not None:
            records.append(
                _ResumeItemRecord(
                    editable=bullet_record.editable,
                    paragraph=paragraph,
                    paragraph_index=paragraph_index,
                    preserved_prefix="",
                )
            )
            continue

        text = paragraph.text.strip()
        if not text:
            continue
        if _is_heading_like(paragraph):
            detected_kind = _section_kind(text)
            level = _heading_level(paragraph)
            if detected_kind is not None or level == 1:
                section_kind = detected_kind or "other"
            continue
        if section_kind not in {"summary", "skills"}:
            continue
        if len(text) > MAX_REPLACEMENT_BULLET_CHARS:
            continue

        item_id = f"{section_kind}-{paragraph_index:04d}"
        preserved_prefix = ""
        editable_text = text
        try:
            _validate_simple_bullet_paragraph(paragraph, item_id)
        except ResumeTailoringError:
            if section_kind != "skills":
                continue
            try:
                preserved_prefix, editable_text = _mixed_skill_prefix(paragraph, item_id)
            except ResumeTailoringError:
                continue
        records.append(
            _ResumeItemRecord(
                editable=EditableBullet(
                    bullet_id=item_id,
                    text=editable_text,
                    section=section_kind,
                    group_id=f"{section_kind}-section",
                ),
                paragraph=paragraph,
                paragraph_index=paragraph_index,
                preserved_prefix=preserved_prefix,
            )
        )

    if not records:
        raise ResumeTailoringError(
            "The resume did not contain safely editable summary, skills, experience, or project "
            "content."
        )
    return tuple(sorted(records, key=lambda record: record.paragraph_index))


def collect_editable_resume_items(document: Any) -> tuple[EditableBullet, ...]:
    """Return all structure-stable text items available to the AI optimizer."""

    return tuple(record.editable for record in _collect_resume_item_records(document))


def _validate_replacement_text(text: str) -> None:
    if not text or text != text.strip():
        raise ResumeTailoringError("Replacement bullet text must be non-empty and trimmed.")
    if "\n" in text or "\r" in text or "\t" in text:
        raise ResumeTailoringError("Replacement bullet text must remain a single paragraph.")
    if len(text) > MAX_REPLACEMENT_BULLET_CHARS:
        raise ResumeTailoringError("Replacement bullet text is too long.")
    if _REPLACEMENT_MARKER.match(text):
        raise ResumeTailoringError("Replacement text must not include a manual bullet marker.")


def _validated_edit_map(
    plan: ValidatedBulletRewritePlan,
    records: tuple[_BulletRecord, ...],
) -> dict[str, tuple[str, ...]]:
    records_by_id = {record.editable.bullet_id: record for record in records}
    if not plan.edits:
        raise ResumeTailoringError("The bullet rewrite plan did not contain any edits.")
    if len(plan.edits) > MAX_EDITED_BULLETS:
        raise ResumeTailoringError("The bullet rewrite plan exceeded the local edit limit.")

    edit_map: dict[str, tuple[str, ...]] = {}
    normalized_replacements: set[str] = set()
    added_bullets = 0
    meaningful_change = False
    for edit in plan.edits:
        if edit.bullet_id in edit_map:
            raise ResumeTailoringError("The bullet rewrite plan targeted a bullet more than once.")
        record = records_by_id.get(edit.bullet_id)
        if record is None:
            raise ResumeTailoringError("The bullet rewrite plan referenced an unknown bullet ID.")
        replacements = edit.replacement_bullets
        if not 1 <= len(replacements) <= 2:
            raise ResumeTailoringError("Each edited bullet must have one or two replacements.")
        for replacement in replacements:
            _validate_replacement_text(replacement)
            normalized = " ".join(replacement.split()).casefold()
            if normalized in normalized_replacements:
                raise ResumeTailoringError(
                    "The bullet rewrite plan contained a duplicate replacement bullet."
                )
            normalized_replacements.add(normalized)
        if edit.bullet_id not in edit.source_bullet_ids:
            raise ResumeTailoringError(
                "Each bullet rewrite must cite its target bullet as source evidence."
            )
        unknown_sources = set(edit.source_bullet_ids).difference(records_by_id)
        if unknown_sources:
            raise ResumeTailoringError(
                "The bullet rewrite plan referenced an unknown source bullet ID."
            )
        if any(
            records_by_id[source_id].editable.group_id != record.editable.group_id
            for source_id in edit.source_bullet_ids
        ):
            raise ResumeTailoringError(
                "The bullet rewrite source bullets must belong to the target bullet's group."
            )
        added_bullets += len(replacements) - 1
        meaningful_change = meaningful_change or replacements != (record.editable.text,)
        edit_map[edit.bullet_id] = replacements

    if added_bullets > MAX_NET_NEW_BULLETS:
        raise ResumeTailoringError("The bullet rewrite plan exceeded the local insertion limit.")
    if not meaningful_change:
        raise ResumeTailoringError("The bullet rewrite plan did not change the resume.")
    return edit_map


def _replace_paragraph_text(paragraph: Paragraph, replacement: str) -> None:
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        raise ResumeTailoringError("A targeted bullet no longer has editable text.")
    text_runs[0].text = replacement
    for run in text_runs[1:]:
        run.text = ""


def _replace_paragraph_text_after_prefix(
    paragraph: Paragraph,
    preserved_prefix: str,
    replacement: str,
) -> None:
    if not preserved_prefix:
        _replace_paragraph_text(paragraph, replacement)
        return
    text_runs = [run for run in paragraph.runs if run.text]
    original = "".join(run.text for run in text_runs)
    if not original.startswith(preserved_prefix):
        raise ResumeTailoringError("A labelled skill paragraph changed before optimization.")

    prefix_length = len(preserved_prefix)
    consumed = 0
    target_index: int | None = None
    prefix_fragment = ""
    for index, run in enumerate(text_runs):
        next_consumed = consumed + len(run.text)
        if prefix_length < next_consumed:
            target_index = index
            prefix_fragment = run.text[: prefix_length - consumed]
            break
        if prefix_length == next_consumed:
            target_index = index + 1
            break
        consumed = next_consumed
    if target_index is None or target_index >= len(text_runs):
        raise ResumeTailoringError("A labelled skill paragraph has no editable content run.")
    text_runs[target_index].text = prefix_fragment + replacement
    for run in text_runs[target_index + 1 :]:
        run.text = ""


def _package_parts(path: Path) -> dict[str, bytes]:
    try:
        with ZipFile(path) as archive:
            # Some Word producers include explicit ZIP directory entries while
            # python-docx omits them on save. Directories carry no document data,
            # so compare only actual OPC package parts.
            return {
                member.filename: archive.read(member)
                for member in archive.infolist()
                if not member.is_dir()
            }
    except (BadZipFile, OSError) as exc:
        raise ResumeTailoringError("Could not inspect the rewritten DOCX package.") from exc


def _canonical_xml(element: Any) -> bytes:
    try:
        return etree.tostring(element, method="c14n")
    except etree.C14NError:
        # Older Office metadata can contain UUID-style namespace names without
        # a URI scheme. C14N 1.0 rejects those names, while C14N 2.0 still gives
        # us a deterministic semantic comparison.
        return etree.tostring(element, method="c14n2")


def _paragraph_text(element: Any) -> str:
    return "".join(node.text or "" for node in element.iter(_W_T))


def _paragraph_structure(element: Any) -> bytes:
    clone = deepcopy(element)
    for text_element in tuple(clone.iter(_W_T)):
        parent = text_element.getparent()
        if parent is not None:
            parent.remove(text_element)
    return _canonical_xml(clone)


def _parse_document_xml(blob: bytes) -> Any:
    try:
        return etree.fromstring(
            blob,
            parser=etree.XMLParser(resolve_entities=False, no_network=True),
        )
    except etree.XMLSyntaxError as exc:
        raise ResumeTailoringError("Could not parse the rewritten DOCX document XML.") from exc


def _unordered_registry_xml_equal(source_root: Any, output_root: Any) -> bool:
    """Compare OPC registry children whose XML order has no semantic meaning."""

    if source_root.tag != output_root.tag or source_root.attrib != output_root.attrib:
        return False

    def child_signatures(root: Any) -> list[bytes]:
        signatures: list[bytes] = []
        for child in root:
            clone = deepcopy(child)
            clone.tail = None
            signatures.append(_canonical_xml(clone))
        return sorted(signatures)

    return child_signatures(source_root) == child_signatures(output_root)


def _package_part_equal(name: str, source_blob: bytes, output_blob: bytes) -> bool:
    if source_blob == output_blob:
        return True
    if not (name.endswith((".xml", ".rels")) or name == "[Content_Types].xml"):
        return False
    try:
        parser = etree.XMLParser(resolve_entities=False, no_network=True)
        source_root = etree.fromstring(source_blob, parser=parser)
        output_root = etree.fromstring(
            output_blob,
            parser=etree.XMLParser(resolve_entities=False, no_network=True),
        )
    except etree.XMLSyntaxError:
        return source_blob == output_blob
    if name == "[Content_Types].xml" or name.endswith(".rels"):
        return _unordered_registry_xml_equal(source_root, output_root)
    return _canonical_xml(source_root) == _canonical_xml(output_root)


def _validated_resume_optimization_map(
    plan: ValidatedBulletRewritePlan,
    records: tuple[_ResumeItemRecord, ...],
) -> dict[str, str]:
    records_by_id = {record.editable.bullet_id: record for record in records}
    if not plan.edits:
        raise ResumeTailoringError("The resume optimization plan did not contain any edits.")
    if len(plan.edits) > MAX_EDITED_BULLETS:
        raise ResumeTailoringError("The resume optimization plan exceeded the local edit limit.")

    edit_map: dict[str, str] = {}
    normalized_replacements: set[str] = set()
    for edit in plan.edits:
        if edit.bullet_id in edit_map:
            raise ResumeTailoringError(
                "The resume optimization plan targeted a paragraph more than once."
            )
        record = records_by_id.get(edit.bullet_id)
        if record is None:
            raise ResumeTailoringError(
                "The resume optimization plan referenced an unknown paragraph ID."
            )
        if len(edit.replacement_bullets) != 1:
            raise ResumeTailoringError(
                "Structure-preserving optimization requires exactly one replacement paragraph."
            )
        replacement = edit.replacement_bullets[0]
        _validate_replacement_text(replacement)
        normalized = " ".join(replacement.split()).casefold()
        if normalized in normalized_replacements:
            raise ResumeTailoringError(
                "The resume optimization plan contained duplicate replacement text."
            )
        if normalized == " ".join(record.editable.text.split()).casefold():
            raise ResumeTailoringError("The resume optimization plan contained a no-op edit.")
        normalized_replacements.add(normalized)
        edit_map[edit.bullet_id] = replacement
    return edit_map


def validate_optimized_resume_docx(
    source_path: Path,
    output_path: Path,
    plan: ValidatedBulletRewritePlan,
) -> None:
    """Prove that only authorized paragraph text changed and structure stayed identical."""

    from docx import Document

    source = validate_resume_path(source_path, tailored=True)
    output = validate_resume_path(output_path, tailored=True)
    try:
        source_document = Document(str(source))
        records = _collect_resume_item_records(source_document)
    except ResumeTailoringError:
        raise
    except Exception as exc:
        raise ResumeTailoringError("Could not inspect the source resume content.") from exc
    edit_map = _validated_resume_optimization_map(plan, records)

    source_parts = _package_parts(source)
    output_parts = _package_parts(output)
    if source_parts.keys() != output_parts.keys():
        raise ResumeTailoringError("DOCX package parts changed during resume optimization.")
    for name, source_blob in source_parts.items():
        if name != DOCUMENT_PART and not _package_part_equal(name, source_blob, output_parts[name]):
            raise ResumeTailoringError(
                "A non-document DOCX part changed during resume optimization."
            )

    source_root = _parse_document_xml(source_parts[DOCUMENT_PART])
    output_root = _parse_document_xml(output_parts[DOCUMENT_PART])
    source_paragraphs = tuple(source_root.iter(_W_P))
    output_paragraphs = tuple(output_root.iter(_W_P))
    if len(output_paragraphs) != len(source_paragraphs):
        raise ResumeTailoringError(
            "Resume optimization added or removed paragraphs and changed the base structure."
        )

    records_by_index = {record.paragraph_index: record for record in records}
    authorized_nodes: list[tuple[Any, Any]] = []
    for paragraph_index, (source_paragraph, output_paragraph) in enumerate(
        zip(source_paragraphs, output_paragraphs, strict=True)
    ):
        record = records_by_index.get(paragraph_index)
        if record is None:
            if _canonical_xml(source_paragraph) != _canonical_xml(output_paragraph):
                raise ResumeTailoringError(
                    "Non-target resume content or formatting changed during optimization."
                )
            continue
        replacement = edit_map.get(record.editable.bullet_id)
        if replacement is None:
            if _canonical_xml(source_paragraph) != _canonical_xml(output_paragraph):
                raise ResumeTailoringError(
                    "Non-target resume content or formatting changed during optimization."
                )
            continue
        if _paragraph_structure(output_paragraph) != _paragraph_structure(source_paragraph):
            raise ResumeTailoringError(
                "A targeted resume paragraph's structure or formatting changed."
            )
        if _paragraph_text(output_paragraph) != f"{record.preserved_prefix}{replacement}":
            raise ResumeTailoringError(
                "A targeted resume paragraph did not contain its validated text."
            )
        authorized_nodes.append((source_paragraph, output_paragraph))

    for source_paragraph, output_paragraph in reversed(authorized_nodes):
        parent = output_paragraph.getparent()
        if parent is None:
            raise ResumeTailoringError("An optimized resume paragraph became detached.")
        parent.replace(output_paragraph, deepcopy(source_paragraph))
    if _canonical_xml(source_root) != _canonical_xml(output_root):
        raise ResumeTailoringError(
            "DOCX structure changed outside the authorized resume text edits."
        )


def create_optimized_resume_docx(
    source_path: Path,
    output_path: Path,
    plan: ValidatedBulletRewritePlan,
) -> Path:
    """Create an atomic, in-place-text-only optimized copy of the base resume."""

    from docx import Document

    source = validate_resume_path(source_path, tailored=True)
    destination = output_path.expanduser()
    if destination.suffix.casefold() != ".docx":
        raise ResumeTailoringError("AI-optimized resume output must be a DOCX file.")
    if destination.resolve(strict=False) == source:
        raise ResumeTailoringError("The source resume must never be overwritten.")

    try:
        document = Document(str(source))
        records = _collect_resume_item_records(document)
        edit_map = _validated_resume_optimization_map(plan, records)
    except ResumeTailoringError:
        raise
    except Exception as exc:
        raise ResumeTailoringError("Could not prepare the source resume for optimization.") from exc

    for record in records:
        replacement = edit_map.get(record.editable.bullet_id)
        if replacement is not None:
            _replace_paragraph_text_after_prefix(
                record.paragraph,
                record.preserved_prefix,
                replacement,
            )

    destination.parent.mkdir(mode=0o700, parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.{uuid4().hex}.tmp.docx")
    try:
        document.save(str(temporary))
        if os.name != "nt":
            os.chmod(temporary, 0o600)
        validate_optimized_resume_docx(source, temporary, plan)
        os.replace(temporary, destination)
        if os.name != "nt":
            os.chmod(destination, 0o600)
    except ResumeTailoringError:
        temporary.unlink(missing_ok=True)
        raise
    except Exception as exc:
        temporary.unlink(missing_ok=True)
        raise ResumeTailoringError("Could not create the optimized resume.") from exc
    return destination


def validate_bullet_rewritten_docx(
    source_path: Path,
    output_path: Path,
    plan: ValidatedBulletRewritePlan,
) -> None:
    """Verify a fresh or cached output contains only the authorized bullet edits."""

    from docx import Document

    source = validate_resume_path(source_path, tailored=True)
    output = validate_resume_path(output_path, tailored=True)
    try:
        source_document = Document(str(source))
        records = _collect_bullet_records(source_document)
    except ResumeTailoringError:
        raise
    except Exception as exc:
        raise ResumeTailoringError("Could not inspect the source resume bullets.") from exc
    edit_map = _validated_edit_map(plan, records)

    source_parts = _package_parts(source)
    output_parts = _package_parts(output)
    if source_parts.keys() != output_parts.keys():
        raise ResumeTailoringError("DOCX package parts changed during bullet rewriting.")
    for name, source_blob in source_parts.items():
        if name != DOCUMENT_PART and not _package_part_equal(name, source_blob, output_parts[name]):
            raise ResumeTailoringError("A non-document DOCX part changed during bullet rewriting.")

    source_root = _parse_document_xml(source_parts[DOCUMENT_PART])
    output_root = _parse_document_xml(output_parts[DOCUMENT_PART])
    source_paragraphs = tuple(source_root.iter(_W_P))
    output_paragraphs = tuple(output_root.iter(_W_P))
    records_by_index = {record.paragraph_index: record for record in records}
    expected_added = sum(len(replacements) - 1 for replacements in edit_map.values())
    if len(output_paragraphs) != len(source_paragraphs) + expected_added:
        raise ResumeTailoringError("Unexpected paragraphs were added or removed during rewriting.")

    output_index = 0
    authorized_nodes: list[tuple[Any, Any, Any | None]] = []
    for source_index, source_paragraph in enumerate(source_paragraphs):
        record = records_by_index.get(source_index)
        replacements = edit_map.get(record.editable.bullet_id) if record is not None else None
        if replacements is None:
            if _canonical_xml(source_paragraph) != _canonical_xml(output_paragraphs[output_index]):
                raise ResumeTailoringError(
                    "Non-target resume content or formatting changed during bullet rewriting."
                )
            output_index += 1
            continue

        first_output = output_paragraphs[output_index]
        if _paragraph_structure(first_output) != _paragraph_structure(source_paragraph):
            raise ResumeTailoringError("A targeted bullet's structure or formatting changed.")
        if _paragraph_text(first_output) != replacements[0]:
            raise ResumeTailoringError("A targeted bullet did not contain its validated text.")
        output_index += 1

        clone_output = None
        if len(replacements) == 2:
            clone_output = output_paragraphs[output_index]
            if (
                clone_output.getparent() is not first_output.getparent()
                or clone_output.getprevious() is not first_output
            ):
                raise ResumeTailoringError(
                    "An added bullet was not cloned immediately after its source bullet."
                )
            if _paragraph_structure(clone_output) != _paragraph_structure(source_paragraph):
                raise ResumeTailoringError("An added bullet did not preserve source formatting.")
            if _paragraph_text(clone_output) != replacements[1]:
                raise ResumeTailoringError("An added bullet did not contain its validated text.")
            output_index += 1
        authorized_nodes.append((source_paragraph, first_output, clone_output))

    if output_index != len(output_paragraphs):
        raise ResumeTailoringError("Unexpected trailing paragraphs were added during rewriting.")

    for source_paragraph, first_output, clone_output in reversed(authorized_nodes):
        if clone_output is not None:
            clone_parent = clone_output.getparent()
            if clone_parent is None:
                raise ResumeTailoringError("The added bullet was detached unexpectedly.")
            clone_parent.remove(clone_output)
        first_parent = first_output.getparent()
        if first_parent is None:
            raise ResumeTailoringError("The rewritten bullet was detached unexpectedly.")
        first_parent.replace(first_output, deepcopy(source_paragraph))

    if _canonical_xml(source_root) != _canonical_xml(output_root):
        raise ResumeTailoringError("DOCX structure changed outside the authorized bullet edits.")


def create_bullet_rewritten_docx(
    source_path: Path,
    output_path: Path,
    plan: ValidatedBulletRewritePlan,
) -> Path:
    """Create an atomic DOCX copy containing only validated bullet replacements."""

    from docx import Document

    source = validate_resume_path(source_path, tailored=True)
    destination = output_path.expanduser()
    if destination.suffix.casefold() != ".docx":
        raise ResumeTailoringError("AI-rewritten resume output must be a DOCX file.")
    if destination.resolve(strict=False) == source:
        raise ResumeTailoringError("The source resume must never be overwritten.")

    try:
        document = Document(str(source))
        records = _collect_bullet_records(document)
        edit_map = _validated_edit_map(plan, records)
    except ResumeTailoringError:
        raise
    except Exception as exc:
        raise ResumeTailoringError(
            "Could not prepare the source resume for bullet rewriting."
        ) from exc

    records_by_id = {record.editable.bullet_id: record for record in records}
    for edit in plan.edits:
        record = records_by_id[edit.bullet_id]
        replacements = edit_map[edit.bullet_id]
        _replace_paragraph_text(record.paragraph, replacements[0])
        if len(replacements) == 2:
            clone_element = deepcopy(record.paragraph._p)
            record.paragraph._p.addnext(clone_element)
            clone_paragraph = Paragraph(clone_element, record.paragraph._parent)
            _replace_paragraph_text(clone_paragraph, replacements[1])

    destination.parent.mkdir(mode=0o700, parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.{uuid4().hex}.tmp.docx")
    try:
        document.save(str(temporary))
        if os.name != "nt":
            os.chmod(temporary, 0o600)
        validate_bullet_rewritten_docx(source, temporary, plan)
        os.replace(temporary, destination)
        if os.name != "nt":
            os.chmod(destination, 0o600)
    except ResumeTailoringError:
        temporary.unlink(missing_ok=True)
        raise
    except Exception as exc:
        temporary.unlink(missing_ok=True)
        raise ResumeTailoringError("Could not create the bullet-rewritten resume.") from exc
    return destination
