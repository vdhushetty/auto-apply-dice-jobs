"""Application service for selecting and preparing a resume per job."""

from __future__ import annotations

import hashlib
import json
import os
import re
from collections import Counter
from collections.abc import Callable, Mapping
from pathlib import Path
from typing import Any

from .bullet_curator import (
    BULLET_REWRITE_PROMPT_VERSION,
    DEFAULT_BULLET_REWRITE_MODEL,
    MAX_EDITED_BULLETS,
    MAX_REPLACEMENT_BULLET_CHARS,
    MAX_REPLACEMENTS_PER_EDIT,
    MAX_SOURCE_BULLETS_PER_EDIT,
    BulletRewritePlanner,
    ValidatedBulletEdit,
    ValidatedBulletRewritePlan,
    validate_bullet_rewrite_plan,
)
from .bullet_documents import (
    collect_editable_bullets,
    create_bullet_rewritten_docx,
    validate_bullet_rewritten_docx,
)
from .curator import (
    PROMPT_VERSION,
    CurationPlanner,
    OpenAIBulletRewritePlanner,
    OpenAICurationPlanner,
    validate_curation_plan,
)
from .documents import (
    SkillSlot,
    collect_skill_slots,
    create_curated_docx,
    extract_resume_text,
    fingerprint_document,
    validate_resume_path,
)
from .layout import DocxLayoutVerifier
from .models import (
    AIReviewPolicy,
    CloudProfile,
    CustomProfile,
    JobPosting,
    MatchDecision,
    PreparedResume,
    ResumeConfigurationError,
    ResumeMode,
    ResumePreparation,
    ResumeTailoringError,
    ResumeVariant,
)
from .selector import (
    DEFAULT_MINIMUM_WINNER_MARGIN,
    ResumeSelector,
    SingleResumeSelector,
    extract_lexical_tokens,
    extract_technology_terms,
)

DEFAULT_THRESHOLD = 35.0
DEFAULT_OUTPUT_DIR = Path(".data/tailored_resumes")
DEFAULT_AI_OUTPUT_DIR = Path(".data/ai_resumes")
_RETRYABLE_AI_BULLET_VALIDATION_ERRORS = frozenset(
    {
        "Bullet rewrite introduced a technology absent from its cited source bullets.",
        "Bullet rewrite edit is a no-op.",
    }
)
_VERIFY_ONLY_AI_FALLBACK_REASONS = frozenset(
    {
        "Bullet rewrite edit is a no-op.",
        "No safe bullet rewrite plan: insufficient_source_evidence.",
        "No safe bullet rewrite plan: no_relevant_change.",
    }
)
LayoutVerifier = Callable[[Path, Path], None]
ReviewCallback = Callable[[JobPosting, Path, str], bool]


class ResumeService:
    """Fail-closed facade used by the Dice browser adapter."""

    def __init__(
        self,
        *,
        mode: ResumeMode,
        paths: Mapping[CloudProfile, str | Path],
        threshold: float = DEFAULT_THRESHOLD,
        minimum_winner_margin: float = DEFAULT_MINIMUM_WINNER_MARGIN,
        output_dir: str | Path = DEFAULT_OUTPUT_DIR,
        ai_resume_path: str | Path = "",
        ai_output_dir: str | Path = DEFAULT_AI_OUTPUT_DIR,
        planner: CurationPlanner | None = None,
        bullet_planner: BulletRewritePlanner | None = None,
        api_key: str | None = None,
        model: str | None = None,
        safety_identity: str = "",
        layout_verifier: LayoutVerifier | None = None,
        ai_review_policy: AIReviewPolicy = AIReviewPolicy.REVIEW_BEFORE_APPLY,
        review_callback: ReviewCallback | None = None,
    ) -> None:
        if not isinstance(ai_review_policy, AIReviewPolicy):
            raise ResumeConfigurationError("AI review policy must be a supported typed value.")
        self.mode = mode
        self._output_dir = (
            Path(ai_output_dir if mode is ResumeMode.AI_BULLETS else output_dir)
            .expanduser()
            .resolve()
        )
        self._planner: CurationPlanner | None = None
        self._bullet_planner: BulletRewritePlanner | None = None
        self._api_key = api_key.strip() if api_key else ""
        self._planner_model: str | None = None
        self._safety_identity = safety_identity
        self._layout_verifier: LayoutVerifier | None = None
        self._ai_review_policy = ai_review_policy
        self._review_callback = review_callback
        self._variants: tuple[ResumeVariant, ...]
        self._selector: ResumeSelector | SingleResumeSelector

        if mode is ResumeMode.AI_BULLETS:
            if not ai_resume_path:
                raise ResumeConfigurationError("A base DOCX resume is required for AI mode.")
            path = validate_resume_path(ai_resume_path, tailored=True)
            text = extract_resume_text(path)
            custom_variant = ResumeVariant(
                profile=CustomProfile.CUSTOM,
                path=path,
                text=text,
                terms=extract_technology_terms(text),
                lexical_tokens=extract_lexical_tokens(text),
            )
            self._variants = (custom_variant,)
            self._selector = SingleResumeSelector(
                custom_variant,
                threshold,
            )
            self._bullet_planner = bullet_planner
            if bullet_planner is not None:
                self._api_key = ""
            self._planner_model = (
                model or os.getenv("OPENAI_MODEL") or DEFAULT_BULLET_REWRITE_MODEL
            ).strip()
            self._layout_verifier = (
                layout_verifier if layout_verifier is not None else DocxLayoutVerifier()
            )
            return

        variants: list[ResumeVariant] = []
        resolved_paths: set[Path] = set()
        content_digests: set[str] = set()
        for profile in CloudProfile:
            path_value = paths.get(profile)
            if not path_value:
                raise ResumeConfigurationError(
                    f"A {profile.value.upper()} resume file is required."
                )
            path = validate_resume_path(path_value, tailored=mode is ResumeMode.TAILORED)
            resolved_paths.add(path)
            content_digests.add(hashlib.sha256(path.read_bytes()).hexdigest())
            text = extract_resume_text(path)
            variants.append(
                ResumeVariant(
                    profile=profile,
                    path=path,
                    text=text,
                    terms=extract_technology_terms(text),
                    lexical_tokens=extract_lexical_tokens(text),
                )
            )
        if len(resolved_paths) != len(CloudProfile):
            raise ResumeConfigurationError("AWS, Azure, and GCP must use three different files.")
        if len(content_digests) != len(CloudProfile):
            raise ResumeConfigurationError(
                "AWS, Azure, and GCP resumes must have distinct file contents."
            )
        self._variants = tuple(variants)
        self._selector = ResumeSelector(
            self._variants,
            threshold,
            minimum_winner_margin=minimum_winner_margin,
        )
        if mode is ResumeMode.TAILORED:
            self._planner = planner
            self._planner_model = model
            self._layout_verifier = (
                layout_verifier if layout_verifier is not None else DocxLayoutVerifier()
            )

    @classmethod
    def from_settings(
        cls,
        settings: Mapping[str, Any],
        *,
        planner: CurationPlanner | None = None,
        bullet_planner: BulletRewritePlanner | None = None,
        api_key: str | None = None,
        safety_identity: str = "",
        layout_verifier: LayoutVerifier | None = None,
        review_callback: ReviewCallback | None = None,
    ) -> ResumeService:
        mode = ResumeMode.parse(str(settings.get("resume_mode", ResumeMode.STATIC.value)))
        raw_paths = settings.get("resume_paths", {})
        if not isinstance(raw_paths, Mapping):
            raise ResumeConfigurationError("resume_paths must be an object.")
        paths = {profile: str(raw_paths.get(profile.value, "")) for profile in CloudProfile}
        threshold = float(settings.get("minimum_match_score", DEFAULT_THRESHOLD))
        minimum_winner_margin = float(
            settings.get("minimum_winner_margin", DEFAULT_MINIMUM_WINNER_MARGIN)
        )
        output_dir = str(settings.get("tailored_resume_output_dir", DEFAULT_OUTPUT_DIR))
        ai_resume_path = str(settings.get("ai_resume_path", "")).strip()
        ai_output_dir = str(settings.get("ai_resume_output_dir", DEFAULT_AI_OUTPUT_DIR))
        model = (
            os.getenv("OPENAI_MODEL", "").strip()
            or str(settings.get("openai_model", "")).strip()
            or None
        )
        raw_review_policy = settings.get(
            "ai_review_policy",
            AIReviewPolicy.REVIEW_BEFORE_APPLY.value,
        )
        if not isinstance(raw_review_policy, str):
            raise ResumeConfigurationError("ai_review_policy must be a string.")
        ai_review_policy = AIReviewPolicy.parse(raw_review_policy)
        return cls(
            mode=mode,
            paths=paths,
            threshold=threshold,
            minimum_winner_margin=minimum_winner_margin,
            output_dir=output_dir,
            ai_resume_path=ai_resume_path,
            ai_output_dir=ai_output_dir,
            planner=planner,
            bullet_planner=bullet_planner,
            api_key=api_key,
            model=model,
            safety_identity=safety_identity,
            layout_verifier=layout_verifier,
            ai_review_policy=ai_review_policy,
            review_callback=review_callback,
        )

    def evaluate(self, job: JobPosting) -> ResumePreparation:
        """Evaluate local fit without generating or modifying a resume."""

        decision = self._selector.select(job)
        return ResumePreparation(
            eligible=decision.eligible,
            reason=decision.reason,
            decision=decision,
        )

    def prepare_selected(self, job: JobPosting, decision: MatchDecision) -> ResumePreparation:
        """Prepare a previously evaluated decision without re-running selection."""

        approved_paths = {variant.path for variant in self._variants}
        if decision.selected_path not in approved_paths:
            return ResumePreparation(
                eligible=False,
                reason="Resume decision referenced an unapproved source file.",
                decision=decision,
            )
        if not decision.eligible:
            return ResumePreparation(
                eligible=False,
                reason=decision.reason,
                decision=decision,
            )
        if self.mode is ResumeMode.STATIC:
            return ResumePreparation(
                eligible=True,
                reason=decision.reason,
                decision=decision,
                prepared=PreparedResume(
                    path=decision.selected_path,
                    decision=decision,
                    tailored=False,
                ),
            )

        try:
            prepared_path = (
                self._prepare_ai_bullets(job, decision.selected_path)
                if self.mode is ResumeMode.AI_BULLETS
                else self._prepare_tailored(job, decision.selected_path)
            )
        except ResumeTailoringError as exc:
            return ResumePreparation(
                eligible=False,
                reason=str(exc),
                decision=decision,
            )
        if (
            self.mode is ResumeMode.AI_BULLETS
            and self._ai_review_policy is AIReviewPolicy.REVIEW_BEFORE_APPLY
        ):
            try:
                self._require_review_approval(job, decision.selected_path, prepared_path)
            except ResumeTailoringError as exc:
                return ResumePreparation(
                    eligible=False,
                    reason=str(exc),
                    decision=decision,
                )
        return ResumePreparation(
            eligible=True,
            reason=decision.reason,
            decision=decision,
            prepared=PreparedResume(
                path=prepared_path,
                decision=decision,
                tailored=True,
            ),
        )

    def prepare(self, job: JobPosting) -> ResumePreparation:
        """Select, gate, and optionally curate a resume for one job."""

        evaluation = self.evaluate(job)
        if not evaluation.eligible or evaluation.decision is None:
            return evaluation
        return self.prepare_selected(job, evaluation.decision)

    def prepare_selected_for_verification(
        self,
        job: JobPosting,
        decision: MatchDecision,
    ) -> ResumePreparation:
        """Prepare one resume for the no-submit Dice upload check.

        A verification run may exercise Dice's file picker with the user-approved source
        DOCX when bullet curation has no safe, non-identical change to make.  This is not
        a curation fallback for submission: normal runs retain the fail-closed result.
        """

        preparation = self.prepare_selected(job, decision)
        if (
            preparation.eligible
            or self.mode is not ResumeMode.AI_BULLETS
            or preparation.reason not in _VERIFY_ONLY_AI_FALLBACK_REASONS
            or not decision.eligible
        ):
            return preparation
        return ResumePreparation(
            eligible=True,
            reason=(
                "Verify-only fallback: AI found no safe non-identical bullet edit; using "
                "the approved base resume to verify Dice upload only."
            ),
            decision=decision,
            prepared=PreparedResume(
                path=decision.selected_path,
                decision=decision,
                tailored=False,
                verification_fallback=True,
            ),
        )

    def assert_prepared_resume_ready(
        self,
        job: JobPosting,
        prepared: PreparedResume,
    ) -> str:
        """Revalidate an AI artifact and enforce its configured review policy before use."""

        try:
            output_digest = self._file_digest(prepared.path)
        except OSError as exc:
            raise ResumeTailoringError("Prepared resume is no longer readable.") from exc
        if self.mode is not ResumeMode.AI_BULLETS:
            return output_digest
        source_path = prepared.decision.selected_path
        if prepared.verification_fallback:
            if prepared.tailored or prepared.path.resolve() != source_path.resolve():
                raise ResumeTailoringError(
                    "Verification fallback must use the exact approved base resume."
                )
            try:
                validate_resume_path(source_path, tailored=True)
            except ResumeConfigurationError as exc:
                raise ResumeTailoringError(
                    "Verification fallback base resume is no longer valid."
                ) from exc
            return output_digest
        if not self._cached_ai_output_is_valid(job, source_path, prepared.path):
            raise ResumeTailoringError(
                "AI-tailored resume changed after validation; the application was skipped."
            )
        if (
            self._ai_review_policy is AIReviewPolicy.REVIEW_BEFORE_APPLY
            and not self._ai_approval_is_valid(job, source_path, prepared.path)
        ):
            raise ResumeTailoringError(
                "AI-tailored resume approval is missing or no longer matches the exact file."
            )
        return output_digest

    def _prepare_tailored(self, job: JobPosting, source_path: Path) -> Path:
        from docx import Document

        if self._planner is None:
            self._planner = OpenAICurationPlanner(
                model=self._planner_model,
                safety_identity=self._safety_identity,
            )
        document = Document(str(source_path))
        slots = collect_skill_slots(document)
        if not slots:
            raise ResumeTailoringError(
                "No safely editable skill list was found in the selected DOCX."
            )
        output_path = self._output_path(job, source_path, self._planner.model)
        if output_path.exists() and self._cached_output_is_valid(source_path, output_path):
            return output_path
        output_path.unlink(missing_ok=True)
        self._manifest_path(output_path).unlink(missing_ok=True)
        raw_plan = self._planner.plan(job, slots)
        plan = validate_curation_plan(raw_plan, job, slots)
        changed_orders = {
            slot_id: order
            for slot_id, order in plan.slot_orders.items()
            if order != next(slot for slot in slots if slot.slot_id == slot_id).original_order
        }
        candidates = self._candidate_orders(slots, changed_orders)

        last_layout_error: ResumeTailoringError | None = None
        for candidate_orders in candidates:
            try:
                create_curated_docx(source_path, output_path, candidate_orders)
                if self._layout_verifier is None:
                    raise ResumeTailoringError("Tailored mode has no layout verifier.")
                self._layout_verifier(source_path, output_path)
                self._write_cache_manifest(source_path, output_path)
                return output_path
            except ResumeTailoringError as exc:
                output_path.unlink(missing_ok=True)
                self._manifest_path(output_path).unlink(missing_ok=True)
                last_layout_error = exc
                if "changed the rendered page count" not in str(exc):
                    raise
        if last_layout_error is not None:
            raise last_layout_error
        raise ResumeTailoringError("No safe tailored resume could be generated.")

    def _prepare_ai_bullets(self, job: JobPosting, source_path: Path) -> Path:
        from docx import Document

        document = Document(str(source_path))
        bullets = collect_editable_bullets(document)
        if not bullets:
            raise ResumeTailoringError(
                "No safely editable experience bullets were found in the base DOCX."
            )
        if self._bullet_planner is None:
            try:
                self._bullet_planner = OpenAIBulletRewritePlanner(
                    api_key=self._api_key or None,
                    model=self._planner_model,
                    safety_identity=self._safety_identity,
                )
            finally:
                self._api_key = ""
        self._planner_model = self._bullet_planner.model
        output_path = self._ai_output_path(job, source_path, self._bullet_planner.model)
        if output_path.exists() and self._cached_ai_output_is_valid(
            job,
            source_path,
            output_path,
        ):
            return output_path

        self._remove_ai_artifacts(output_path)
        try:
            raw_plan = self._bullet_planner.plan(job, bullets)
        except ResumeTailoringError:
            raise
        except Exception as exc:
            raise ResumeTailoringError("OpenAI could not produce a bullet rewrite plan.") from exc
        source_text = next(
            variant.text for variant in self._variants if variant.path == source_path
        )
        try:
            plan = validate_bullet_rewrite_plan(raw_plan, job, bullets, source_text)
        except ResumeTailoringError as first_error:
            retry_plan = getattr(self._bullet_planner, "retry_plan", None)
            if str(first_error) not in _RETRYABLE_AI_BULLET_VALIDATION_ERRORS or not callable(
                retry_plan
            ):
                raise
            try:
                retry_raw_plan = retry_plan(job, bullets)
                plan = validate_bullet_rewrite_plan(
                    retry_raw_plan,
                    job,
                    bullets,
                    source_text,
                )
            except ResumeTailoringError as retry_error:
                raise ResumeTailoringError(str(retry_error)) from retry_error
            except Exception as retry_error:
                raise ResumeTailoringError(
                    "OpenAI could not produce a conservative bullet rewrite retry."
                ) from retry_error
        try:
            create_bullet_rewritten_docx(source_path, output_path, plan)
            validate_bullet_rewritten_docx(source_path, output_path, plan)
            if self._layout_verifier is None:
                raise ResumeTailoringError("AI bullet mode has no layout verifier.")
            self._layout_verifier(source_path, output_path)
            self._write_ai_cache_manifest(job, source_path, output_path, plan)
            return output_path
        except ResumeTailoringError:
            self._remove_ai_artifacts(output_path)
            raise
        except Exception as exc:
            self._remove_ai_artifacts(output_path)
            raise ResumeTailoringError("AI bullet resume generation failed safely.") from exc

    @staticmethod
    def _serialize_bullet_plan(plan: ValidatedBulletRewritePlan) -> dict[str, Any]:
        return {
            "reason_code": plan.reason_code,
            "edits": [
                {
                    "bullet_id": edit.bullet_id,
                    "replacement_bullets": list(edit.replacement_bullets),
                    "source_bullet_ids": list(edit.source_bullet_ids),
                }
                for edit in plan.edits
            ],
        }

    @staticmethod
    def _deserialize_bullet_plan(value: object) -> ValidatedBulletRewritePlan:
        if not isinstance(value, dict):
            raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
        reason_code = value.get("reason_code")
        raw_edits = value.get("edits")
        if not isinstance(reason_code, str) or not isinstance(raw_edits, list) or not raw_edits:
            raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
        if len(raw_edits) > MAX_EDITED_BULLETS:
            raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
        edits: list[ValidatedBulletEdit] = []
        seen_bullet_ids: set[str] = set()
        seen_replacements: set[str] = set()
        for raw_edit in raw_edits:
            if not isinstance(raw_edit, dict):
                raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
            bullet_id = raw_edit.get("bullet_id")
            replacements = raw_edit.get("replacement_bullets")
            source_ids = raw_edit.get("source_bullet_ids")
            if (
                not isinstance(bullet_id, str)
                or bullet_id in seen_bullet_ids
                or not isinstance(replacements, list)
                or not replacements
                or len(replacements) > MAX_REPLACEMENTS_PER_EDIT
                or not all(isinstance(item, str) and item for item in replacements)
                or any(len(item) > MAX_REPLACEMENT_BULLET_CHARS for item in replacements)
                or not isinstance(source_ids, list)
                or not source_ids
                or len(source_ids) > MAX_SOURCE_BULLETS_PER_EDIT
                or not all(isinstance(item, str) and item for item in source_ids)
                or len(set(source_ids)) != len(source_ids)
                or bullet_id not in source_ids
            ):
                raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
            normalized_replacements = {" ".join(item.split()).casefold() for item in replacements}
            if len(normalized_replacements) != len(replacements) or seen_replacements.intersection(
                normalized_replacements
            ):
                raise ResumeTailoringError("AI resume cache contained an invalid rewrite plan.")
            seen_bullet_ids.add(bullet_id)
            seen_replacements.update(normalized_replacements)
            edits.append(
                ValidatedBulletEdit(
                    bullet_id=bullet_id,
                    replacement_bullets=tuple(replacements),
                    source_bullet_ids=tuple(source_ids),
                )
            )
        return ValidatedBulletRewritePlan(edits=tuple(edits), reason_code=reason_code)

    def _cached_ai_output_is_valid(
        self,
        job: JobPosting,
        source_path: Path,
        output_path: Path,
    ) -> bool:
        manifest_path = self._manifest_path(output_path)
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            if not isinstance(manifest, dict):
                return False
            if manifest.get("mode") != ResumeMode.AI_BULLETS.value:
                return False
            if manifest.get("layout_verified") is not True:
                return False
            if manifest.get("prompt_version") != BULLET_REWRITE_PROMPT_VERSION:
                return False
            if manifest.get("model") != self._planner_model:
                return False
            if manifest.get("job_sha256") != self._job_digest(job):
                return False
            if manifest.get("source_sha256") != self._file_digest(source_path):
                return False
            if manifest.get("output_sha256") != self._file_digest(output_path):
                return False
            plan = self._deserialize_bullet_plan(manifest.get("validated_plan"))
            validate_resume_path(output_path, tailored=True)
            validate_bullet_rewritten_docx(source_path, output_path, plan)
            return True
        except Exception:
            return False

    def _write_ai_cache_manifest(
        self,
        job: JobPosting,
        source_path: Path,
        output_path: Path,
        plan: ValidatedBulletRewritePlan,
    ) -> None:
        manifest_path = self._manifest_path(output_path)
        manifest_path.parent.mkdir(mode=0o700, parents=True, exist_ok=True)
        temporary_path = manifest_path.with_suffix(f"{manifest_path.suffix}.tmp")
        payload = {
            "mode": ResumeMode.AI_BULLETS.value,
            "layout_verified": True,
            "source_sha256": self._file_digest(source_path),
            "output_sha256": self._file_digest(output_path),
            "job_sha256": self._job_digest(job),
            "prompt_version": BULLET_REWRITE_PROMPT_VERSION,
            "model": self._bullet_planner.model if self._bullet_planner is not None else "",
            "validated_plan": self._serialize_bullet_plan(plan),
        }
        temporary_path.write_text(json.dumps(payload, sort_keys=True), encoding="utf-8")
        if temporary_path.stat().st_size == 0:
            raise ResumeTailoringError("AI resume cache manifest could not be written.")
        if os.name != "nt":
            os.chmod(temporary_path, 0o600)
        temporary_path.replace(manifest_path)

    def _require_review_approval(
        self,
        job: JobPosting,
        source_path: Path,
        output_path: Path,
    ) -> None:
        if self._ai_approval_is_valid(job, source_path, output_path):
            return
        output_digest = self._file_digest(output_path)
        if self._review_callback is None:
            raise ResumeTailoringError(
                "AI-tailored resume requires explicit review approval before submission."
            )
        try:
            approved = self._review_callback(job, output_path, output_digest)
        except Exception as exc:
            raise ResumeTailoringError("AI resume review could not be completed safely.") from exc
        if not approved:
            raise ResumeTailoringError(
                "AI-tailored resume was not approved; the application was skipped."
            )
        try:
            current_digest = self._file_digest(output_path)
        except OSError as exc:
            raise ResumeTailoringError(
                "AI-tailored resume became unreadable during review."
            ) from exc
        if current_digest != output_digest:
            raise ResumeTailoringError(
                "AI-tailored resume changed during review; regenerate and review it again."
            )
        try:
            self._write_ai_approval(job, source_path, output_path, output_digest)
        except Exception as exc:
            raise ResumeTailoringError("AI resume approval could not be recorded safely.") from exc

    def _ai_approval_is_valid(
        self,
        job: JobPosting,
        source_path: Path,
        output_path: Path,
    ) -> bool:
        approval_path = self._approval_path(output_path)
        manifest_path = self._manifest_path(output_path)
        try:
            approval = json.loads(approval_path.read_text(encoding="utf-8"))
            return bool(
                isinstance(approval, dict)
                and approval.get("approved") is True
                and approval.get("source_sha256") == self._file_digest(source_path)
                and approval.get("output_sha256") == self._file_digest(output_path)
                and approval.get("job_sha256") == self._job_digest(job)
                and approval.get("manifest_sha256") == self._file_digest(manifest_path)
            )
        except Exception:
            return False

    def _write_ai_approval(
        self,
        job: JobPosting,
        source_path: Path,
        output_path: Path,
        expected_output_digest: str,
    ) -> None:
        approval_path = self._approval_path(output_path)
        temporary_path = approval_path.with_suffix(f"{approval_path.suffix}.tmp")
        if self._file_digest(output_path) != expected_output_digest:
            raise ResumeTailoringError("AI-tailored resume changed before approval was recorded.")
        payload = {
            "approved": True,
            "source_sha256": self._file_digest(source_path),
            "output_sha256": expected_output_digest,
            "job_sha256": self._job_digest(job),
            "manifest_sha256": self._file_digest(self._manifest_path(output_path)),
        }
        temporary_path.write_text(json.dumps(payload, sort_keys=True), encoding="utf-8")
        if os.name != "nt":
            os.chmod(temporary_path, 0o600)
        temporary_path.replace(approval_path)

    @staticmethod
    def _approval_path(output_path: Path) -> Path:
        return output_path.with_suffix(f"{output_path.suffix}.approval.json")

    def _remove_ai_artifacts(self, output_path: Path) -> None:
        output_path.unlink(missing_ok=True)
        self._manifest_path(output_path).unlink(missing_ok=True)
        self._approval_path(output_path).unlink(missing_ok=True)

    @staticmethod
    def _job_digest(job: JobPosting) -> str:
        value = f"{job.url}\n{job.title}\n{job.description}"
        return hashlib.sha256(value.encode()).hexdigest()

    def _ai_output_path(self, job: JobPosting, source_path: Path, model: str) -> Path:
        title_slug = re.sub(r"[^a-z0-9]+", "-", job.title.lower()).strip("-")[:60]
        source_stat = source_path.stat()
        cache_material = "|".join(
            (
                job.url,
                hashlib.sha256(job.description.encode()).hexdigest(),
                str(source_path),
                str(source_stat.st_mtime_ns),
                str(source_stat.st_size),
                BULLET_REWRITE_PROMPT_VERSION,
                model,
            )
        )
        digest = hashlib.sha256(cache_material.encode()).hexdigest()[:12]
        return self._output_dir / f"{title_slug or 'dice-job'}-{digest}.docx"

    @staticmethod
    def _candidate_orders(
        slots: tuple[SkillSlot, ...],
        changed_orders: dict[str, tuple[str, ...]],
    ) -> tuple[dict[str, tuple[str, ...]], ...]:
        """Try the model plan, then smaller model-consistent page-safe edits."""

        originals = {slot.slot_id: slot.original_order for slot in slots}
        candidates: list[dict[str, tuple[str, ...]]] = []
        seen: set[tuple[tuple[str, tuple[str, ...]], ...]] = set()

        def add(candidate: dict[str, tuple[str, ...]]) -> None:
            key = tuple(sorted(candidate.items()))
            if candidate and key not in seen:
                candidates.append(candidate)
                seen.add(key)

        add(changed_orders)
        for slot_id, target_order in changed_orders.items():
            add({slot_id: target_order})
            original_order = originals[slot_id]
            target_positions = {item_id: index for index, item_id in enumerate(target_order)}
            for index in range(len(original_order) - 1):
                left = original_order[index]
                right = original_order[index + 1]
                if target_positions[left] <= target_positions[right]:
                    continue
                adjacent_order = list(original_order)
                adjacent_order[index], adjacent_order[index + 1] = right, left
                add({slot_id: tuple(adjacent_order)})
        return tuple(candidates)

    @staticmethod
    def _manifest_path(output_path: Path) -> Path:
        return output_path.with_suffix(f"{output_path.suffix}.manifest.json")

    @staticmethod
    def _file_digest(path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    def _cached_output_is_valid(self, source_path: Path, output_path: Path) -> bool:
        from docx import Document

        manifest_path = self._manifest_path(output_path)
        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            if not isinstance(manifest, dict) or manifest.get("layout_verified") is not True:
                return False
            if manifest.get("source_sha256") != self._file_digest(source_path):
                return False
            if manifest.get("output_sha256") != self._file_digest(output_path):
                return False
            validate_resume_path(output_path, tailored=True)
            source_document = Document(str(source_path))
            cached_document = Document(str(output_path))
            if fingerprint_document(cached_document) != fingerprint_document(source_document):
                return False
            return Counter(extract_resume_text(source_path)) == Counter(
                extract_resume_text(output_path)
            )
        except Exception:
            return False

    def _write_cache_manifest(self, source_path: Path, output_path: Path) -> None:
        manifest_path = self._manifest_path(output_path)
        manifest_path.parent.mkdir(mode=0o700, parents=True, exist_ok=True)
        temporary_path = manifest_path.with_suffix(f"{manifest_path.suffix}.tmp")
        payload = {
            "layout_verified": True,
            "source_sha256": self._file_digest(source_path),
            "output_sha256": self._file_digest(output_path),
            "prompt_version": PROMPT_VERSION,
            "model": self._planner.model if self._planner is not None else "",
        }
        temporary_path.write_text(json.dumps(payload, sort_keys=True), encoding="utf-8")
        if temporary_path.stat().st_size == 0:
            raise ResumeTailoringError("Tailored resume cache manifest could not be written.")
        if os.name != "nt":
            os.chmod(temporary_path, 0o600)
        temporary_path.replace(manifest_path)

    def _output_path(self, job: JobPosting, source_path: Path, model: str) -> Path:
        title_slug = re.sub(r"[^a-z0-9]+", "-", job.title.lower()).strip("-")[:60]
        source_stat = source_path.stat()
        cache_material = "|".join(
            (
                job.url,
                hashlib.sha256(job.description.encode()).hexdigest(),
                str(source_path),
                str(source_stat.st_mtime_ns),
                str(source_stat.st_size),
                PROMPT_VERSION,
                model,
            )
        )
        digest = hashlib.sha256(cache_material.encode()).hexdigest()[:12]
        return self._output_dir / f"{title_slug or 'dice-job'}-{digest}.docx"
