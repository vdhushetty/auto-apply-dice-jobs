"""Resume text extraction and structure-preserving DOCX edits."""

from __future__ import annotations

import hashlib
import os
import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from .models import ResumeConfigurationError, ResumeTailoringError
from .selector import extract_technology_terms

SUPPORTED_STATIC_SUFFIXES = frozenset({".docx", ".pdf"})
SUPPORTED_TAILORED_SUFFIXES = frozenset({".docx"})
MAX_RESUME_BYTES = 10 * 1024 * 1024
SKILL_SECTION_WORDS = frozenset(
    {
        "competencies",
        "skills",
        "technologies",
        "technology",
        "technical skills",
        "technical competencies",
        "tools",
    }
)
SKILL_CATEGORY_LABELS = frozenset(
    {
        "analytics",
        "cloud",
        "databases",
        "frameworks",
        "languages",
        "machine learning",
        "platforms",
        "programming languages",
        "tools",
    }
)
ACTION_VERBS = frozenset(
    {
        "architected",
        "automated",
        "built",
        "created",
        "delivered",
        "deployed",
        "designed",
        "developed",
        "implemented",
        "improved",
        "led",
        "maintained",
        "managed",
        "migrated",
        "optimized",
        "reduced",
        "supported",
    }
)
DELIMITER_PATTERN = re.compile(r"([,;|\u2022]\s*)")
WORDPROCESSINGML_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
VML_NAMESPACE = "urn:schemas-microsoft-com:vml"

TRACKED_CHANGE_TAGS = frozenset(
    {
        f"{{{WORDPROCESSINGML_NAMESPACE}}}cellDel",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}cellIns",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}cellMerge",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}del",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}ins",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveFrom",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveFromRangeEnd",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveFromRangeStart",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveTo",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveToRangeEnd",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}moveToRangeStart",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}numberingChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}pPrChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}rPrChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}sectPrChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}tblPrChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}tcPrChange",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}trPrChange",
    }
)
CONTENT_CONTROL_TAG = f"{{{WORDPROCESSINGML_NAMESPACE}}}sdt"
TEXT_BOX_TAGS = frozenset(
    {
        f"{{{WORDPROCESSINGML_NAMESPACE}}}txbxContent",
        f"{{{VML_NAMESPACE}}}textbox",
    }
)
ALT_CHUNK_TAG = f"{{{WORDPROCESSINGML_NAMESPACE}}}altChunk"
DRAWING_TAG = f"{{{WORDPROCESSINGML_NAMESPACE}}}drawing"
OBJECT_TAGS = frozenset(
    {
        f"{{{WORDPROCESSINGML_NAMESPACE}}}object",
        f"{{{WORDPROCESSINGML_NAMESPACE}}}pict",
    }
)


@dataclass(frozen=True)
class SkillItem:
    item_id: str
    text: str


@dataclass
class SkillSlot:
    slot_id: str
    paragraph: Any
    prefix: str
    items: tuple[SkillItem, ...]
    separators: tuple[str, ...]
    body_run_indexes: tuple[int, ...]
    suffix: str = ""

    @property
    def original_order(self) -> tuple[str, ...]:
        return tuple(item.item_id for item in self.items)

    def render(self, ordered_item_ids: tuple[str, ...]) -> str:
        by_id = {item.item_id: item.text for item in self.items}
        values = [by_id[item_id] for item_id in ordered_item_ids]
        body_parts: list[str] = []
        for index, value in enumerate(values):
            body_parts.append(value)
            if index < len(self.separators):
                body_parts.append(self.separators[index])
        return self.prefix + "".join(body_parts) + self.suffix


@dataclass(frozen=True)
class DocumentFingerprint:
    section_geometry: tuple[tuple[int | None, ...], ...]
    paragraph_styles: tuple[str, ...]
    table_geometry: tuple[tuple[int, tuple[int, ...]], ...]
    header_footer_hash: str
    paragraph_count: int
    inline_shape_count: int
    embedded_media_hashes: tuple[str, ...]
    drawing_count: int
    object_count: int


def _audit_tailored_docx(path: Path) -> None:
    """Reject OOXML constructs that cannot be safely preserved during tailoring."""

    found_tags: set[str] = set()
    try:
        with ZipFile(path) as archive:
            for member in archive.infolist():
                if not member.filename.startswith("word/") or not member.filename.endswith(".xml"):
                    continue
                try:
                    root = ET.fromstring(archive.read(member))
                except ET.ParseError as exc:
                    raise ResumeConfigurationError(
                        "Tailored mode could not inspect the DOCX XML safely."
                    ) from exc
                found_tags.update(
                    element.tag for element in root.iter() if isinstance(element.tag, str)
                )
    except (BadZipFile, OSError) as exc:
        raise ResumeConfigurationError(
            "Tailored mode requires a valid, readable DOCX package."
        ) from exc

    if found_tags.intersection(TRACKED_CHANGE_TAGS):
        raise ResumeConfigurationError(
            "Tailored mode does not support DOCX files with tracked changes."
        )
    if CONTENT_CONTROL_TAG in found_tags:
        raise ResumeConfigurationError(
            "Tailored mode does not support DOCX files with content controls."
        )
    if found_tags.intersection(TEXT_BOX_TAGS):
        raise ResumeConfigurationError("Tailored mode does not support DOCX files with text boxes.")
    if ALT_CHUNK_TAG in found_tags:
        raise ResumeConfigurationError(
            "Tailored mode does not support DOCX files with altChunk content."
        )


def validate_resume_path(path_value: str | Path, *, tailored: bool = False) -> Path:
    """Resolve and validate a local resume path without following unsafe inputs."""

    path = Path(path_value).expanduser()
    try:
        resolved = path.resolve(strict=True)
    except (FileNotFoundError, OSError) as exc:
        raise ResumeConfigurationError(f"Resume file does not exist: {path}") from exc
    if not resolved.is_file():
        raise ResumeConfigurationError(f"Resume path is not a file: {resolved}")
    if path.is_symlink():
        raise ResumeConfigurationError(f"Symbolic-link resumes are not accepted: {path}")
    allowed = SUPPORTED_TAILORED_SUFFIXES if tailored else SUPPORTED_STATIC_SUFFIXES
    if resolved.suffix.lower() not in allowed:
        suffixes = ", ".join(sorted(allowed))
        raise ResumeConfigurationError(
            f"Unsupported resume format '{resolved.suffix}'. Allowed: {suffixes}."
        )
    if resolved.stat().st_size > MAX_RESUME_BYTES:
        raise ResumeConfigurationError("Resume files must be 10 MB or smaller.")
    if tailored:
        _audit_tailored_docx(resolved)
    return resolved


def extract_resume_text(path: Path) -> str:
    """Extract text used for matching; source content is never logged."""

    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text for paragraph in _iter_all_paragraphs(document)]
            text = "\n".join(paragraphs)
        elif suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            raise ResumeConfigurationError(f"Unsupported resume format: {suffix}")
    except ResumeConfigurationError:
        raise
    except Exception as exc:
        raise ResumeConfigurationError(f"Could not read resume file: {path.name}") from exc
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized) < 100:
        raise ResumeConfigurationError(
            f"Resume '{path.name}' did not contain enough extractable text."
        )
    return normalized


def _iter_table_paragraphs(table: Any) -> list[Any]:
    paragraphs: list[Any] = []
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_identity = id(cell._tc)
            if cell_identity in seen_cells:
                continue
            seen_cells.add(cell_identity)
            paragraphs.extend(cell.paragraphs)
            for nested_table in cell.tables:
                paragraphs.extend(_iter_table_paragraphs(nested_table))
    return paragraphs


def _iter_all_paragraphs(document: Any) -> list[Any]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(_iter_table_paragraphs(table))
    for section in document.sections:
        for container in (section.header, section.footer):
            paragraphs.extend(container.paragraphs)
            for table in container.tables:
                paragraphs.extend(_iter_table_paragraphs(table))
    unique: list[Any] = []
    seen: set[int] = set()
    for paragraph in paragraphs:
        identity = id(paragraph._p)
        if identity not in seen:
            unique.append(paragraph)
            seen.add(identity)
    return unique


def _normalized_heading(text: str) -> str:
    return re.sub(r"[^a-z ]+", " ", text.lower()).strip()


def _is_heading(paragraph: Any) -> bool:
    text = paragraph.text.strip()
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if style_name.startswith(("heading", "title")):
        return True
    if not text or len(text) > 80:
        return False
    normalized = _normalized_heading(text)
    return text.isupper() or normalized in SKILL_SECTION_WORDS


def _run_format_signature(run: Any) -> str:
    properties = run._r.rPr
    return str(properties.xml) if properties is not None else ""


def _body_run_indexes(paragraph: Any, body_offset: int) -> tuple[int, ...]:
    indexes: list[int] = []
    cursor = 0
    for index, run in enumerate(paragraph.runs):
        next_cursor = cursor + len(run.text)
        if next_cursor > body_offset and run.text:
            indexes.append(index)
        cursor = next_cursor
    if not indexes:
        return ()
    signatures = {
        _run_format_signature(paragraph.runs[index])
        for index in indexes
        if paragraph.runs[index].text
    }
    return tuple(indexes) if len(signatures) <= 1 else ()


def _parse_skill_slot(paragraph: Any, slot_index: int, in_skill_section: bool) -> SkillSlot | None:
    text = paragraph.text.strip()
    if not text:
        return None
    lower = text.lower()
    prefix_label = _normalized_heading(lower.split(":", 1)[0])
    has_skill_prefix = prefix_label in SKILL_CATEGORY_LABELS or prefix_label in SKILL_SECTION_WORDS
    if not (in_skill_section or has_skill_prefix):
        return None

    leading = paragraph.text[: len(paragraph.text) - len(paragraph.text.lstrip())]
    content = paragraph.text.lstrip()
    prefix = leading
    body = content
    if ":" in content:
        label, candidate_body = content.split(":", 1)
        if len(label) <= 40:
            prefix += f"{label}:"
            body = candidate_body
    body_leading = body[: len(body) - len(body.lstrip())]
    prefix += body_leading
    body = body.lstrip()

    trailing = body[len(body.rstrip()) :]
    body = body.rstrip()
    suffix = trailing
    if body.endswith(".") and not body.endswith(".."):
        body = body[:-1]
        suffix = "." + suffix

    parts = DELIMITER_PATTERN.split(body)
    item_texts = tuple(part.strip() for part in parts[0::2] if part.strip())
    separators = tuple(parts[1::2])
    if len(item_texts) < 3 or len(separators) != len(item_texts) - 1:
        return None
    if len({item.lower() for item in item_texts}) != len(item_texts):
        return None
    if any(
        len(item) > 60
        or len(item.split()) > 6
        or any(mark in item for mark in (".", "?", "!"))
        or item.split()[0].casefold() in ACTION_VERBS
        for item in item_texts
    ):
        return None
    if (
        not in_skill_section
        and sum(bool(extract_technology_terms(item)) for item in item_texts) < 2
    ):
        return None

    body_offset = len(prefix)
    run_indexes = _body_run_indexes(paragraph, body_offset)
    if not run_indexes:
        return None
    slot_id = f"skill-slot-{slot_index:03d}"
    items = tuple(
        SkillItem(item_id=f"{slot_id}-item-{index:03d}", text=item)
        for index, item in enumerate(item_texts)
    )
    return SkillSlot(
        slot_id=slot_id,
        paragraph=paragraph,
        prefix=prefix,
        items=items,
        separators=separators,
        suffix=suffix,
        body_run_indexes=run_indexes,
    )


def collect_skill_slots(document: Any) -> tuple[SkillSlot, ...]:
    """Collect safely editable, delimiter-based lists from skill sections."""

    slots: list[SkillSlot] = []
    in_skill_section = False
    for paragraph in _iter_all_paragraphs(document):
        if _is_heading(paragraph):
            heading = _normalized_heading(paragraph.text)
            in_skill_section = any(word in heading for word in SKILL_SECTION_WORDS)
            continue
        slot = _parse_skill_slot(paragraph, len(slots), in_skill_section)
        if slot is not None:
            slots.append(slot)
    return tuple(slots)


def fingerprint_document(document: Any) -> DocumentFingerprint:
    """Capture layout-sensitive structure that curation is not allowed to alter."""

    geometry: list[tuple[int | None, ...]] = []
    header_footer_parts: list[str] = []
    for section in document.sections:
        geometry.append(
            (
                section.page_width,
                section.page_height,
                section.top_margin,
                section.right_margin,
                section.bottom_margin,
                section.left_margin,
            )
        )
        header_footer_parts.extend([section.header._element.xml, section.footer._element.xml])
    table_geometry = tuple(
        (len(table.rows), tuple(len(row.cells) for row in table.rows)) for table in document.tables
    )
    paragraphs = _iter_all_paragraphs(document)
    styles = tuple(paragraph.style.name if paragraph.style else "" for paragraph in paragraphs)
    media_hashes: list[str] = []
    drawing_count = 0
    object_count = 0
    for part in document.part.package.parts:
        part_name = str(part.partname)
        blob = part.blob
        if part_name.startswith("/word/media/"):
            media_hashes.append(hashlib.sha256(blob).hexdigest())
        if not part.content_type.endswith(("xml", "+xml")):
            continue
        try:
            root = ET.fromstring(blob)
        except ET.ParseError:
            continue
        for element in root.iter():
            if element.tag == DRAWING_TAG:
                drawing_count += 1
            elif element.tag in OBJECT_TAGS:
                object_count += 1
    return DocumentFingerprint(
        section_geometry=tuple(geometry),
        paragraph_styles=styles,
        table_geometry=table_geometry,
        header_footer_hash=hashlib.sha256("".join(header_footer_parts).encode()).hexdigest(),
        paragraph_count=len(paragraphs),
        inline_shape_count=len(document.inline_shapes),
        embedded_media_hashes=tuple(sorted(media_hashes)),
        drawing_count=drawing_count,
        object_count=object_count,
    )


def _replace_slot_order(slot: SkillSlot, ordered_item_ids: tuple[str, ...]) -> None:
    original_ids = slot.original_order
    if len(ordered_item_ids) != len(original_ids) or set(ordered_item_ids) != set(original_ids):
        raise ResumeTailoringError(f"Plan for {slot.slot_id} is not an exact item permutation.")
    rendered = slot.render(ordered_item_ids)
    body = rendered[len(slot.prefix) :]
    runs = slot.paragraph.runs
    first_index = slot.body_run_indexes[0]
    cursor = 0
    prefix_remainder = ""
    for index, run in enumerate(runs):
        next_cursor = cursor + len(run.text)
        if index == first_index:
            local_prefix_length = max(0, len(slot.prefix) - cursor)
            prefix_remainder = run.text[:local_prefix_length]
            break
        cursor = next_cursor
    runs[first_index].text = prefix_remainder + body
    for index in slot.body_run_indexes[1:]:
        runs[index].text = ""


def create_curated_docx(
    source_path: Path,
    output_path: Path,
    slot_orders: dict[str, tuple[str, ...]],
) -> Path:
    """Copy a DOCX and reorder exact candidate-authored skill items only."""

    from docx import Document

    document = Document(str(source_path))
    before = fingerprint_document(document)
    slots = {slot.slot_id: slot for slot in collect_skill_slots(document)}
    if not slot_orders:
        raise ResumeTailoringError("The curation plan did not contain any safe changes.")
    unknown_slots = set(slot_orders).difference(slots)
    if unknown_slots:
        raise ResumeTailoringError("The curation plan referenced an unknown skill slot.")
    changed = False
    for slot_id, order in slot_orders.items():
        slot = slots[slot_id]
        if order != slot.original_order:
            _replace_slot_order(slot, order)
            changed = True
    if not changed:
        raise ResumeTailoringError("The curation plan did not change the resume.")

    output_path.parent.mkdir(mode=0o700, parents=True, exist_ok=True)
    document.save(str(output_path))
    if os.name != "nt":
        os.chmod(output_path, 0o600)
    reloaded = Document(str(output_path))
    after = fingerprint_document(reloaded)
    if before != after:
        output_path.unlink(missing_ok=True)
        raise ResumeTailoringError("DOCX structure changed during curation; output was discarded.")
    return output_path
