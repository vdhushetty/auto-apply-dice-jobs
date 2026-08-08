from __future__ import annotations

import hashlib
import json
from collections.abc import Sequence
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from core.resumes.bullet_curator import EditableBullet
from core.resumes.documents import SkillSlot
from core.resumes.models import (
    AIReviewPolicy,
    JobPosting,
    ResumeConfigurationError,
    ResumeTailoringError,
)
from core.resumes.service import ResumeService


class ReversePlanner:
    model = "fake-model"

    def plan(self, job: JobPosting, slots: tuple[SkillSlot, ...]) -> dict[str, Any]:
        slot = slots[0]
        return {
            "schema_version": "1",
            "outcome": "curate",
            "reason_code": "ok",
            "job_evidence": [{"quote": "AWS", "priority": "required"}],
            "slot_orders": [
                {
                    "slot_id": slot.slot_id,
                    "ordered_item_ids": list(reversed(slot.original_order)),
                }
            ],
        }


class ExplodingPlanner:
    model = "fake-model"

    def plan(self, job: JobPosting, slots: tuple[SkillSlot, ...]) -> dict[str, Any]:
        raise RuntimeError("provider unavailable")


class CountingReversePlanner(ReversePlanner):
    def __init__(self) -> None:
        self.calls = 0

    def plan(self, job: JobPosting, slots: tuple[SkillSlot, ...]) -> dict[str, Any]:
        self.calls += 1
        return super().plan(job, slots)


class AIBulletPlanner:
    model = "fake-ai-model"

    def __init__(self) -> None:
        self.calls = 0

    def plan(self, job: JobPosting, bullets: Sequence[EditableBullet]) -> dict[str, Any]:
        self.calls += 1
        target = next(item for item in bullets if item.section == "experience")
        return {
            "schema_version": "1",
            "outcome": "rewrite",
            "reason_code": "ok",
            "job_evidence": [{"quote": "AWS", "priority": "required"}],
            "edits": [
                {
                    "bullet_id": target.bullet_id,
                    "replacement_bullets": [
                        "Built reliable AWS data pipelines with Python and SQL for governed "
                        "business reporting."
                    ],
                    "source_bullet_ids": [target.bullet_id],
                }
            ],
        }


class RetryingAIBulletPlanner(AIBulletPlanner):
    def __init__(self) -> None:
        super().__init__()
        self.retry_calls = 0

    def plan(self, job: JobPosting, bullets: Sequence[EditableBullet]) -> dict[str, Any]:
        self.calls += 1
        target = [item for item in bullets if item.section == "experience"][1]
        return {
            "schema_version": "1",
            "outcome": "rewrite",
            "reason_code": "ok",
            "job_evidence": [{"quote": "AWS", "priority": "required"}],
            "edits": [
                {
                    "bullet_id": target.bullet_id,
                    "replacement_bullets": [
                        "Automated AWS data quality checks and monitored production batch workflows."
                    ],
                    "source_bullet_ids": [target.bullet_id],
                }
            ],
        }

    def retry_plan(
        self,
        job: JobPosting,
        bullets: Sequence[EditableBullet],
    ) -> dict[str, Any]:
        self.retry_calls += 1
        return super().plan(job, bullets)


class LayoutRetryPlanner(AIBulletPlanner):
    def __init__(self) -> None:
        super().__init__()
        self.retry_calls = 0

    def retry_plan(
        self,
        job: JobPosting,
        bullets: Sequence[EditableBullet],
    ) -> dict[str, Any]:
        self.retry_calls += 1
        return super().plan(job, bullets)


class NoOpAIBulletPlanner:
    model = "fake-ai-model"

    def plan(self, job: JobPosting, bullets: Sequence[EditableBullet]) -> dict[str, Any]:
        target = bullets[0]
        return {
            "schema_version": "1",
            "outcome": "rewrite",
            "reason_code": "ok",
            "job_evidence": [{"quote": "AWS", "priority": "required"}],
            "edits": [
                {
                    "bullet_id": target.bullet_id,
                    "replacement_bullets": [target.text],
                    "source_bullet_ids": [target.bullet_id],
                }
            ],
        }


class ExplodingAIBulletPlanner:
    model = "fake-ai-model"

    def plan(self, job: JobPosting, bullets: Sequence[EditableBullet]) -> dict[str, Any]:
        raise AssertionError("AI planner must not run during evaluation")


class WholeResumePlanner:
    model = "fake-whole-resume-model"

    def plan(self, job: JobPosting, items: Sequence[EditableBullet]) -> dict[str, Any]:
        summary = next(item for item in items if item.section == "summary")
        skills = next(item for item in items if item.section == "skills")
        experience = next(
            item for item in items if item.section == "experience" and "AWS" in item.text
        )
        return {
            "schema_version": "1",
            "outcome": "rewrite",
            "reason_code": "ok",
            "job_evidence": [{"quote": "AWS, Python, and SQL", "priority": "required"}],
            "edits": [
                {
                    "bullet_id": summary.bullet_id,
                    "replacement_bullets": [
                        "AWS data engineer building reliable analytics platforms and governed "
                        "reporting systems with Python and SQL for business stakeholders."
                    ],
                    "source_bullet_ids": [summary.bullet_id, experience.bullet_id],
                },
                {
                    "bullet_id": skills.bullet_id,
                    "replacement_bullets": ["AWS, Python, SQL, data quality, batch workflows"],
                    "source_bullet_ids": [skills.bullet_id, experience.bullet_id],
                },
                {
                    "bullet_id": experience.bullet_id,
                    "replacement_bullets": [
                        "Built governed AWS data pipelines with Python and SQL for reliable "
                        "business reporting."
                    ],
                    "source_bullet_ids": [experience.bullet_id],
                },
            ],
        }


def make_ai_resume(path: Path) -> Path:
    document = Document()
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(
        "Data engineer with production experience building reliable analytics platforms and "
        "governed reporting systems for business stakeholders."
    )
    document.add_heading("Technical Skills", level=1)
    document.add_paragraph("Python, SQL, AWS, batch workflows, data quality")
    document.add_heading("Professional Experience", level=1)
    document.add_paragraph("Senior Data Engineer | Example Company")
    document.add_paragraph(
        "Built AWS data pipelines with Python and SQL for reliable business reporting.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Automated data quality checks and monitored production batch workflows.",
        style="List Bullet",
    )
    document.save(path)
    return path


def test_ai_optimizer_updates_summary_skills_and_experience_without_restructuring(
    tmp_path: Path,
) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    source_document = Document(source)
    source_paragraph_count = len(source_document.paragraphs)
    source_styles = [paragraph.style.name for paragraph in source_document.paragraphs]
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 95,
        },
        bullet_planner=WholeResumePlanner(),
        layout_verifier=lambda source_path, output_path: None,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL for governed data pipelines.",
            url="https://www.dice.com/job-detail/whole-resume-optimizer",
        )
    )

    assert result.eligible and result.prepared is not None
    output_document = Document(result.prepared.path)
    assert len(output_document.paragraphs) == source_paragraph_count
    assert [paragraph.style.name for paragraph in output_document.paragraphs] == source_styles
    output_text = "\n".join(paragraph.text for paragraph in output_document.paragraphs)
    assert "AWS data engineer building reliable analytics platforms" in output_text
    assert "AWS, Python, SQL, data quality, batch workflows" in output_text
    assert "Built governed AWS data pipelines" in output_text

    manifest_path = result.prepared.path.with_suffix(f"{result.prepared.path.suffix}.manifest.json")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert len(manifest["validated_plan"]["edits"]) == 3
    assert {"aws", "python", "sql"} <= set(manifest["optimization_audit"]["incorporated_job_terms"])


def test_ai_optimizer_reduces_plan_until_protected_layout_fits(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    layout_calls = 0

    def accept_only_one_changed_paragraph(source_path: Path, output_path: Path) -> None:
        nonlocal layout_calls
        layout_calls += 1
        source_text = [paragraph.text for paragraph in Document(source_path).paragraphs]
        output_text = [paragraph.text for paragraph in Document(output_path).paragraphs]
        changed = sum(left != right for left, right in zip(source_text, output_text, strict=True))
        if changed > 1:
            raise ResumeTailoringError(
                "Tailored resume moved a protected section heading to another rendered page "
                "and was discarded."
            )

    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=WholeResumePlanner(),
        layout_verifier=accept_only_one_changed_paragraph,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL for governed data pipelines.",
            url="https://www.dice.com/job-detail/ai-layout-backoff",
        )
    )

    assert result.eligible and result.prepared is not None
    manifest_path = result.prepared.path.with_suffix(f"{result.prepared.path.suffix}.manifest.json")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert len(manifest["validated_plan"]["edits"]) == 1
    assert layout_calls >= 2


def make_paths(resume_factory) -> dict[str, str]:  # type: ignore[no-untyped-def]
    return {
        "aws": str(resume_factory("aws.docx", "AWS", ["Python", "SQL", "AWS", "S3", "Glue"])),
        "azure": str(
            resume_factory(
                "azure.docx", "Azure", ["Python", "SQL", "Azure", "Data Factory", "Synapse"]
            )
        ),
        "gcp": str(
            resume_factory("gcp.docx", "GCP", ["Python", "SQL", "GCP", "BigQuery", "Dataflow"])
        ),
    }


def test_static_mode_selects_without_calling_ai(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    service = ResumeService.from_settings(
        {
            "resume_mode": "static",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ExplodingPlanner(),
    )

    result = service.prepare(
        JobPosting(
            title="GCP Data Engineer",
            description="Build GCP BigQuery and Dataflow pipelines with Python and SQL.",
            url="https://www.dice.com/job-detail/static",
        )
    )

    assert result.eligible
    assert result.prepared is not None
    assert result.prepared.path.name == "gcp.docx"
    assert not result.prepared.tailored


def test_tailored_mode_creates_job_specific_copy(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    source_paths = make_paths(resume_factory)
    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": source_paths,
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ReversePlanner(),
        layout_verifier=lambda source, output: None,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required AWS Glue and S3 data pipelines using Python and SQL.",
            url="https://www.dice.com/job-detail/tailored",
        )
    )

    assert result.eligible
    assert result.prepared is not None
    assert result.prepared.tailored
    assert result.prepared.path.exists()
    assert result.prepared.path != Path(source_paths["aws"])


def test_evaluate_does_not_generate_or_call_planner(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ExplodingPlanner(),
        layout_verifier=lambda source, output: None,
    )

    result = service.evaluate(
        JobPosting(
            title="AWS Data Engineer",
            description="Build AWS Glue and S3 pipelines using Python and SQL.",
            url="https://www.dice.com/job-detail/evaluate",
        )
    )

    assert result.eligible
    assert result.prepared is None
    assert not (tmp_path / "out").exists()


def test_page_count_drift_discards_tailored_output(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    def reject_layout(source: Path, output: Path) -> None:
        raise ResumeTailoringError(
            "Tailored resume changed the rendered page count and was discarded."
        )

    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ReversePlanner(),
        layout_verifier=reject_layout,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required AWS Glue and S3 data pipelines using Python and SQL.",
            url="https://www.dice.com/job-detail/layout-drift",
        )
    )

    assert not result.eligible
    assert "page count" in result.reason
    assert not list((tmp_path / "out").glob("*.docx"))


def test_layout_drift_retries_a_smaller_model_consistent_reorder(
    resume_factory, tmp_path: Path
) -> None:  # type: ignore[no-untyped-def]
    verifier_calls = 0

    def accept_second_candidate(source: Path, output: Path) -> None:
        nonlocal verifier_calls
        verifier_calls += 1
        if verifier_calls == 1:
            raise ResumeTailoringError(
                "Tailored resume changed the rendered page count and was discarded."
            )

    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ReversePlanner(),
        layout_verifier=accept_second_candidate,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required AWS Glue and S3 data pipelines using Python and SQL.",
            url="https://www.dice.com/job-detail/layout-fallback",
        )
    )

    assert result.eligible
    assert result.prepared is not None
    assert verifier_calls == 2


def test_duplicate_resume_content_is_rejected(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    aws = resume_factory("aws.docx", "AWS", ["Python", "SQL", "AWS", "S3", "Glue"])
    duplicate = tmp_path / "duplicate.docx"
    duplicate.write_bytes(aws.read_bytes())
    gcp = resume_factory("gcp.docx", "GCP", ["Python", "SQL", "GCP", "BigQuery", "Dataflow"])

    with pytest.raises(ResumeConfigurationError, match="distinct file contents"):
        ResumeService.from_settings(
            {
                "resume_mode": "static",
                "resume_paths": {
                    "aws": str(aws),
                    "azure": str(duplicate),
                    "gcp": str(gcp),
                },
            }
        )


def test_tampered_cached_body_is_regenerated(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    planner = CountingReversePlanner()
    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 20,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=planner,
        layout_verifier=lambda source, output: None,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required AWS Glue and S3 data pipelines using Python and SQL.",
        url="https://www.dice.com/job-detail/cache-tamper",
    )
    first = service.prepare(job)
    assert first.prepared is not None
    output = first.prepared.path

    tampered = Document(str(output))
    tampered.paragraphs[0].text = "Tampered body text"
    tampered.save(str(output))
    manifest_path = output.with_suffix(f"{output.suffix}.manifest.json")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["output_sha256"] = hashlib.sha256(output.read_bytes()).hexdigest()
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")

    second = service.prepare(job)

    assert second.eligible
    assert second.prepared is not None
    assert planner.calls == 2
    assert "Tampered body text" not in "\n".join(
        paragraph.text for paragraph in Document(str(second.prepared.path)).paragraphs
    )


def test_low_match_skips_before_curation(resume_factory, tmp_path: Path) -> None:  # type: ignore[no-untyped-def]
    service = ResumeService.from_settings(
        {
            "resume_mode": "tailored",
            "resume_paths": make_paths(resume_factory),
            "minimum_match_score": 90,
            "tailored_resume_output_dir": str(tmp_path / "out"),
        },
        planner=ExplodingPlanner(),
    )

    result = service.prepare(
        JobPosting(
            title="Mainframe Administrator",
            description="RACF, COBOL, z/OS, CICS and mainframe security are required.",
            url="https://www.dice.com/job-detail/skip",
        )
    )

    assert not result.eligible
    assert result.prepared is None


def test_ai_bullets_uses_one_resume_and_hash_bound_review(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    planner = AIBulletPlanner()
    reviews: list[tuple[str, Path, str]] = []

    def approve(job: JobPosting, path: Path, digest: str) -> bool:
        reviews.append((job.title, path, digest))
        return True

    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=planner,
        api_key="sk-test-secret-not-real",
        layout_verifier=lambda source_path, output_path: None,
        review_callback=approve,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-reviewed",
    )

    first = service.prepare(job)
    second = service.prepare(job)

    assert first.eligible and first.prepared is not None
    assert first.prepared.tailored
    assert first.prepared.path != source
    assert second.eligible and second.prepared is not None
    assert second.prepared.path == first.prepared.path
    assert planner.calls == 1
    assert len(reviews) == 1
    assert reviews[0][2] == hashlib.sha256(first.prepared.path.read_bytes()).hexdigest()
    approval = first.prepared.path.with_suffix(f"{first.prepared.path.suffix}.approval.json")
    assert approval.exists()
    artifact_text = "\n".join(
        path.read_text(encoding="utf-8")
        for path in (approval, first.prepared.path.with_suffix(".docx.manifest.json"))
    )
    assert "sk-test-secret-not-real" not in artifact_text


def test_ai_bullets_attempts_curation_despite_low_initial_match(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    planner = AIBulletPlanner()
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 95,
            "ai_review_policy": "skip_review",
        },
        bullet_planner=planner,
        layout_verifier=lambda source_path, output_path: None,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS Glue, S3, Python, SQL, and Terraform.",
        url="https://www.dice.com/job-detail/ai-low-initial-match",
    )

    result = service.prepare(job)

    assert result.eligible and result.prepared is not None
    assert result.decision is not None
    assert result.decision.tailoring_match_gate_bypassed
    assert planner.calls == 1


def test_ai_bullets_retries_one_source_ungrounded_technology_plan(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    planner = RetryingAIBulletPlanner()
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=planner,
        layout_verifier=lambda source_path, output_path: None,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-source-grounded-retry",
    )

    result = service.prepare(job)

    assert result.eligible and result.prepared is not None
    assert planner.calls == 2
    assert planner.retry_calls == 1
    assert source.exists()


def test_ai_optimizer_retries_once_after_protected_layout_drift(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    planner = LayoutRetryPlanner()
    layout_calls = 0

    def reject_first_layout(source_path: Path, output_path: Path) -> None:
        nonlocal layout_calls
        layout_calls += 1
        if layout_calls == 1:
            raise ResumeTailoringError(
                "Tailored resume moved a protected section heading to another rendered page "
                "and was discarded."
            )

    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=planner,
        layout_verifier=reject_first_layout,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL data pipeline experience.",
            url="https://www.dice.com/job-detail/ai-layout-retry",
        )
    )

    assert result.eligible and result.prepared is not None
    assert layout_calls == 2
    assert planner.calls == 2
    assert planner.retry_calls == 1


def test_ai_bullets_verify_upload_can_use_source_after_noop_plan(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=NoOpAIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-verify-noop",
    )
    evaluation = service.evaluate(job)
    assert evaluation.eligible and evaluation.decision is not None

    submission_preparation = service.prepare_selected(job, evaluation.decision)
    assert not submission_preparation.eligible
    assert submission_preparation.reason == "Bullet rewrite edit is a no-op."

    verification_preparation = service.prepare_selected_for_verification(
        job,
        evaluation.decision,
    )

    assert verification_preparation.eligible and verification_preparation.prepared is not None
    assert verification_preparation.prepared.path == source
    assert not verification_preparation.prepared.tailored
    assert verification_preparation.prepared.verification_fallback
    assert (
        service.assert_prepared_resume_ready(job, verification_preparation.prepared)
        == hashlib.sha256(source.read_bytes()).hexdigest()
    )


def test_ai_bullets_review_rejection_skips_upload_candidate(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
        review_callback=lambda job, path, digest: False,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL data pipeline experience.",
            url="https://www.dice.com/job-detail/ai-rejected",
        )
    )

    assert not result.eligible
    assert result.prepared is None
    assert "not approved" in result.reason
    assert not list((tmp_path / "ai-out").glob("*.approval.json"))


def test_ai_bullets_cannot_prepare_without_review_callback(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL data pipeline experience.",
            url="https://www.dice.com/job-detail/ai-review-required",
        )
    )

    assert not result.eligible
    assert result.prepared is None
    assert "requires explicit review approval" in result.reason


def test_ai_bullets_preview_evaluation_needs_no_key_or_planner_call(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "minimum_match_score": 20,
        },
        bullet_planner=ExplodingAIBulletPlanner(),
    )

    result = service.evaluate(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL data pipeline experience.",
            url="https://www.dice.com/job-detail/ai-preview",
        )
    )

    assert result.eligible
    assert result.prepared is None
    assert result.decision is not None
    assert result.decision.selected_profile.value == "custom"


def test_ai_bullets_rejects_file_changed_during_review(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")

    def mutate_during_review(job: JobPosting, path: Path, digest: str) -> bool:
        document = Document(path)
        document.paragraphs[0].text = "Changed while the review dialog was open"
        document.save(path)
        return True

    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
        review_callback=mutate_during_review,
    )

    result = service.prepare(
        JobPosting(
            title="AWS Data Engineer",
            description="Required: AWS, Python, and SQL data pipeline experience.",
            url="https://www.dice.com/job-detail/ai-review-mutation",
        )
    )

    assert not result.eligible
    assert "changed during review" in result.reason
    assert not list((tmp_path / "ai-out").glob("*.approval.json"))


def test_ai_bullets_pre_upload_guard_rejects_post_approval_change(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
        review_callback=lambda job, path, digest: True,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-post-approval-mutation",
    )
    result = service.prepare(job)
    assert result.prepared is not None
    expected_digest = service.assert_prepared_resume_ready(job, result.prepared)
    assert expected_digest == hashlib.sha256(result.prepared.path.read_bytes()).hexdigest()

    tampered = Document(result.prepared.path)
    tampered.paragraphs[0].text = "Changed after approval"
    tampered.save(result.prepared.path)

    with pytest.raises(ResumeTailoringError, match="changed after validation"):
        service.assert_prepared_resume_ready(job, result.prepared)


def test_ai_bullets_skip_review_never_calls_callback_or_requires_approval(
    tmp_path: Path,
) -> None:
    source = make_ai_resume(tmp_path / "base.docx")

    def forbidden_review(job: JobPosting, path: Path, digest: str) -> bool:
        raise AssertionError("skip_review must never invoke the review callback")

    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": "skip_review",
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
        review_callback=forbidden_review,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-skip-review",
    )

    result = service.prepare(job)

    assert result.eligible and result.prepared is not None
    assert (
        service.assert_prepared_resume_ready(job, result.prepared)
        == hashlib.sha256(result.prepared.path.read_bytes()).hexdigest()
    )
    assert not list((tmp_path / "ai-out").glob("*.approval.json"))


def test_ai_bullets_skip_review_guard_still_rejects_tampering(tmp_path: Path) -> None:
    source = make_ai_resume(tmp_path / "base.docx")
    service = ResumeService.from_settings(
        {
            "resume_mode": "ai_bullets",
            "ai_review_policy": AIReviewPolicy.SKIP_REVIEW.value,
            "ai_resume_path": str(source),
            "ai_resume_output_dir": str(tmp_path / "ai-out"),
            "minimum_match_score": 20,
        },
        bullet_planner=AIBulletPlanner(),
        layout_verifier=lambda source_path, output_path: None,
    )
    job = JobPosting(
        title="AWS Data Engineer",
        description="Required: AWS, Python, and SQL data pipeline experience.",
        url="https://www.dice.com/job-detail/ai-skip-review-tamper",
    )
    result = service.prepare(job)
    assert result.prepared is not None

    tampered = Document(result.prepared.path)
    tampered.paragraphs[0].text = "Changed after validation"
    tampered.save(result.prepared.path)

    with pytest.raises(ResumeTailoringError, match="changed after validation"):
        service.assert_prepared_resume_ready(job, result.prepared)


@pytest.mark.parametrize("value", ["unchecked", "", 7, None])
def test_invalid_ai_review_policy_fails_closed(
    resume_factory,
    value: object,
) -> None:  # type: ignore[no-untyped-def]
    with pytest.raises(ResumeConfigurationError, match="review policy|ai_review_policy"):
        ResumeService.from_settings(
            {
                "resume_mode": "static",
                "resume_paths": make_paths(resume_factory),
                "ai_review_policy": value,
            }
        )
