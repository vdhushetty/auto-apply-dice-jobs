from __future__ import annotations

import struct
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from core.resumes.bullet_curator import (
    ValidatedBulletEdit,
    ValidatedBulletRewritePlan,
)
from core.resumes.bullet_documents import (
    _package_part_equal,
    collect_editable_bullets,
    collect_editable_resume_items,
    create_bullet_rewritten_docx,
    create_optimized_resume_docx,
    validate_bullet_rewritten_docx,
    validate_optimized_resume_docx,
)
from core.resumes.models import ResumeTailoringError


def _write_bitmap(path: Path) -> None:
    path.write_bytes(
        struct.pack("<2sIHHI", b"BM", 58, 0, 0, 54)
        + struct.pack("<IiiHHIIiiII", 40, 1, 1, 1, 24, 0, 4, 2835, 2835, 0, 0)
        + bytes((180, 80, 20, 0))
    )


def _add_direct_numbering(paragraph) -> None:  # type: ignore[no-untyped-def]
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), "1")
    numbering.extend((level, number_id))
    paragraph_properties.append(numbering)


def _make_resume(path: Path) -> Path:
    badge = path.with_suffix(".bmp")
    _write_bitmap(badge)
    document = Document()
    document.sections[0].left_margin = 720000
    document.sections[0].right_margin = 720000
    document.sections[0].header.paragraphs[0].text = "Resume header"
    document.sections[0].footer.paragraphs[0].text = "Resume footer"
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(
        "Built reliable platforms across several teams, but this summary is not a bullet."
    )
    document.add_heading("Professional Experience", level=1)
    document.add_paragraph("Senior Data Engineer | Example Company")

    first = document.add_paragraph(style="List Bullet")
    first_run = first.add_run(
        "Built AWS data pipelines with Python and SQL for reliable business reporting."
    )
    first_run.bold = True

    second = document.add_paragraph()
    _add_direct_numbering(second)
    second_run = second.add_run(
        "Automated data quality checks and monitored production batch workflows."
    )
    second_run.italic = True

    document.add_paragraph("Data Engineer | Another Company")
    fallback = document.add_paragraph(
        "Optimized Spark processing workflows while preserving operational controls."
    )
    fallback.paragraph_format.left_indent = Inches(0.25)
    fallback.runs[0].underline = True
    document.add_paragraph(
        "- Built a manually marked line that must not be treated as an editable bullet."
    )

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Non-target table content"
    document.add_picture(str(badge))
    document.save(path)
    return path


def _plan(
    bullet_id: str,
    *replacements: str,
    source_bullet_ids: tuple[str, ...] | None = None,
) -> ValidatedBulletRewritePlan:
    return ValidatedBulletRewritePlan(
        edits=(
            ValidatedBulletEdit(
                bullet_id=bullet_id,
                replacement_bullets=tuple(replacements),
                source_bullet_ids=source_bullet_ids or (bullet_id,),
            ),
        ),
        reason_code="ok",
    )


def test_collects_numbered_and_experience_fallback_bullets_with_stable_opaque_ids(
    tmp_path: Path,
) -> None:
    source = _make_resume(tmp_path / "resume.docx")

    first_collection = collect_editable_bullets(Document(source))
    second_collection = collect_editable_bullets(Document(source))

    assert first_collection == second_collection
    assert [bullet.text for bullet in first_collection] == [
        "Built AWS data pipelines with Python and SQL for reliable business reporting.",
        "Automated data quality checks and monitored production batch workflows.",
        "Optimized Spark processing workflows while preserving operational controls.",
    ]
    assert [bullet.section for bullet in first_collection] == [
        "experience",
        "experience",
        "experience",
    ]
    assert first_collection[0].group_id == first_collection[1].group_id == "group-0001"
    assert first_collection[2].group_id == "group-0002"
    assert all("Company" not in bullet.group_id for bullet in first_collection)


def test_collects_whole_resume_items_without_exposing_headings(tmp_path: Path) -> None:
    source = _make_resume(tmp_path / "resume.docx")
    document = Document(source)
    experience_heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "Professional Experience"
    )
    skills_heading = document.paragraphs[experience_heading_index].insert_paragraph_before(
        "Technical Skills"
    )
    skills_heading.style = document.styles["Heading 1"]
    document.paragraphs[experience_heading_index + 1].insert_paragraph_before(
        "Python, SQL, AWS, Spark, data quality"
    )
    document.save(source)

    items = collect_editable_resume_items(Document(source))

    assert {item.section for item in items} == {"summary", "skills", "experience"}
    assert all("Professional" not in item.text for item in items)
    assert any(item.text.startswith("Built reliable platforms") for item in items)
    assert any(item.text.startswith("Python, SQL, AWS") for item in items)


def test_whole_resume_optimization_preserves_exact_document_structure(tmp_path: Path) -> None:
    source = _make_resume(tmp_path / "source.docx")
    document = Document(source)
    experience_heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "Professional Experience"
    )
    skills_heading = document.paragraphs[experience_heading_index].insert_paragraph_before(
        "Technical Skills"
    )
    skills_heading.style = document.styles["Heading 1"]
    document.paragraphs[experience_heading_index + 1].insert_paragraph_before(
        "Python, SQL, AWS, Spark, data quality"
    )
    document.save(source)
    source_bytes = source.read_bytes()
    source_document = Document(source)
    items = collect_editable_resume_items(source_document)
    summary = next(item for item in items if item.section == "summary")
    skills = next(item for item in items if item.section == "skills")
    experience = next(item for item in items if item.section == "experience" and "AWS" in item.text)
    plan = ValidatedBulletRewritePlan(
        edits=(
            ValidatedBulletEdit(
                summary.bullet_id,
                (
                    "Built reliable AWS data platforms and governed reporting systems across "
                    "several business teams.",
                ),
                (summary.bullet_id, experience.bullet_id),
            ),
            ValidatedBulletEdit(
                skills.bullet_id,
                ("AWS, Python, SQL, Spark, data quality",),
                (skills.bullet_id, experience.bullet_id),
            ),
            ValidatedBulletEdit(
                experience.bullet_id,
                (
                    "Built governed AWS data pipelines with Python and SQL for reliable "
                    "business reporting.",
                ),
                (experience.bullet_id,),
            ),
        ),
        reason_code="ok",
    )

    output = create_optimized_resume_docx(source, tmp_path / "optimized.docx", plan)
    validate_optimized_resume_docx(source, output, plan)

    output_document = Document(output)
    assert source.read_bytes() == source_bytes
    assert len(output_document.paragraphs) == len(source_document.paragraphs)
    assert [p.style.name for p in output_document.paragraphs] == [
        p.style.name for p in source_document.paragraphs
    ]
    assert len(output_document.tables) == len(source_document.tables)
    assert output_document.sections[0].header.paragraphs[0].text == "Resume header"
    assert output_document.sections[0].footer.paragraphs[0].text == "Resume footer"


def test_optimization_preserves_mixed_format_skill_label_runs(tmp_path: Path) -> None:
    source = tmp_path / "labelled-skills.docx"
    document = Document()
    document.add_heading("Technical Skills", level=1)
    paragraph = document.add_paragraph()
    label = paragraph.add_run("Cloud Platforms:")
    label.bold = True
    content = paragraph.add_run(" Azure, AWS, GCP")
    content.bold = False
    document.add_heading("Professional Experience", level=1)
    document.add_paragraph(
        "Built AWS and Azure data pipelines for governed reporting.",
        style="List Bullet",
    )
    document.save(source)
    items = collect_editable_resume_items(Document(source))
    skill = next(item for item in items if item.section == "skills")
    assert skill.text == "Azure, AWS, GCP"
    plan = _plan(skill.bullet_id, "AWS, Azure, GCP")

    output = create_optimized_resume_docx(source, tmp_path / "optimized.docx", plan)
    validate_optimized_resume_docx(source, output, plan)

    optimized = Document(output).paragraphs[1]
    assert optimized.text == "Cloud Platforms: AWS, Azure, GCP"
    assert optimized.runs[0].text == "Cloud Platforms:"
    assert optimized.runs[0].bold is True
    assert optimized.runs[1].text == " AWS, Azure, GCP"
    assert optimized.runs[1].bold is False


def test_unnumbered_role_heading_is_not_treated_as_action_verb_bullet(tmp_path: Path) -> None:
    source = tmp_path / "role-heading.docx"
    document = Document()
    document.add_heading("Professional Experience", level=1)
    role_heading = document.add_paragraph("Managed Services Engineer | Acme Corporation")
    role_heading.paragraph_format.left_indent = Inches(0.25)
    genuine_bullet = document.add_paragraph(
        "Managed production data services while preserving operational controls."
    )
    genuine_bullet.paragraph_format.left_indent = Inches(0.25)
    document.save(source)

    bullets = collect_editable_bullets(Document(source))

    assert [bullet.text for bullet in bullets] == [genuine_bullet.text]


def test_numbered_bullets_in_protected_and_compound_sections_are_not_editable(
    tmp_path: Path,
) -> None:
    source = tmp_path / "protected-sections.docx"
    document = Document()
    for heading, text in (
        ("Professional Summary", "Built a numbered summary statement."),
        ("TECHNICAL SKILLS & TOOLS", "Python, SQL, and AWS"),
        ("EDUCATION & TRAINING", "Completed a numbered degree entry."),
        ("CERTIFICATIONS & LICENSES", "Earned a numbered certification entry."),
    ):
        document.add_heading(heading, level=1)
        document.add_paragraph(text, style="List Bullet")
    document.add_heading("Professional Experience", level=1)
    document.add_paragraph("ACME CORPORATION")
    document.add_paragraph(
        "Built a numbered production data workflow.",
        style="List Bullet",
    )
    document.add_heading("Selected Projects", level=1)
    document.add_paragraph(
        "Created a numbered analytics project workflow.",
        style="List Bullet",
    )
    document.save(source)

    bullets = collect_editable_bullets(Document(source))

    assert [bullet.section for bullet in bullets] == ["experience", "projects"]
    assert [bullet.text for bullet in bullets] == [
        "Built a numbered production data workflow.",
        "Created a numbered analytics project workflow.",
    ]


def test_rewrite_preserves_non_target_content_formatting_numbering_and_package_parts(
    tmp_path: Path,
) -> None:
    source = _make_resume(tmp_path / "source.docx")
    source_document = Document(source)
    bullets = collect_editable_bullets(source_document)
    target = bullets[0]
    source_target = next(
        paragraph for paragraph in source_document.paragraphs if paragraph.text == target.text
    )
    source_properties = source_target._p.pPr.xml
    source_media = {
        member.filename: blob
        for member, blob in _zip_members(source)
        if member.filename.startswith("word/media/")
    }
    replacement = (
        "Built reliable AWS pipelines with Python and SQL for governed business reporting."
    )
    plan = _plan(target.bullet_id, replacement)

    output = create_bullet_rewritten_docx(source, tmp_path / "rewritten.docx", plan)
    validate_bullet_rewritten_docx(source, output, plan)

    output_document = Document(output)
    output_target = next(
        paragraph for paragraph in output_document.paragraphs if paragraph.text == replacement
    )
    assert output_target.style.name == source_target.style.name == "List Bullet"
    assert output_target._p.pPr.xml == source_properties
    assert output_target.runs[0].bold is True
    assert output_document.sections[0].left_margin == source_document.sections[0].left_margin
    assert output_document.sections[0].right_margin == source_document.sections[0].right_margin
    assert output_document.sections[0].header.paragraphs[0].text == "Resume header"
    assert output_document.sections[0].footer.paragraphs[0].text == "Resume footer"
    assert output_document.tables[0].cell(0, 0).text == "Non-target table content"
    output_media = {
        member.filename: blob
        for member, blob in _zip_members(output)
        if member.filename.startswith("word/media/")
    }
    assert output_media == source_media


def test_rewrite_ignores_non_content_zip_directory_entries(tmp_path: Path) -> None:
    source = _make_resume(tmp_path / "source-with-directories.docx")
    with ZipFile(source, mode="a") as archive:
        existing = set(archive.namelist())
        for directory in ("word/", "word/media/", "docProps/"):
            if directory not in existing:
                archive.writestr(directory, b"")

    target = collect_editable_bullets(Document(source))[0]
    plan = _plan(
        target.bullet_id,
        "Built reliable AWS pipelines with Python and SQL for governed reporting.",
    )

    output = create_bullet_rewritten_docx(source, tmp_path / "rewritten.docx", plan)

    validate_bullet_rewritten_docx(source, output, plan)


def test_office_uuid_namespace_metadata_uses_safe_canonical_fallback() -> None:
    source = b'<root xmlns:x="ba0ffe40-8f88-40d6-a542-c6c8ce1c4353"><x:child value="same"/></root>'
    reformatted = (
        b'<root xmlns:x="ba0ffe40-8f88-40d6-a542-c6c8ce1c4353">'
        b'<x:child value="same"></x:child></root>'
    )

    assert _package_part_equal("customXml/item.xml", source, reformatted)


def test_content_type_registry_order_is_not_treated_as_document_drift() -> None:
    source = (
        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        b'<Default Extension="xml" ContentType="application/xml"/>'
        b'<Override PartName="/word/document.xml" '
        b'ContentType="application/document"/></Types>'
    )
    reordered = (
        b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        b'<Override ContentType="application/document" '
        b'PartName="/word/document.xml"/>'
        b'<Default ContentType="application/xml" Extension="xml"/></Types>'
    )

    assert _package_part_equal("[Content_Types].xml", source, reordered)


def _zip_members(path: Path):  # type: ignore[no-untyped-def]
    with ZipFile(path) as archive:
        return [(member, archive.read(member)) for member in archive.infolist()]


def test_structure_preserving_optimizer_rejects_paragraph_splits(
    tmp_path: Path,
) -> None:
    source = _make_resume(tmp_path / "source.docx")
    source_document = Document(source)
    target = collect_editable_bullets(source_document)[1]
    source_target = next(
        paragraph for paragraph in source_document.paragraphs if paragraph.text == target.text
    )
    source_properties = source_target._p.pPr.xml
    first_replacement = "Automated Python data quality checks for production batch workflows."
    second_replacement = "Monitored SQL batch workflows using existing operational controls."
    plan = _plan(target.bullet_id, first_replacement, second_replacement)

    with pytest.raises(ResumeTailoringError, match="insertion limit"):
        create_bullet_rewritten_docx(source, tmp_path / "split.docx", plan)

    assert source_properties


def test_unknown_target_fails_closed_without_creating_output(tmp_path: Path) -> None:
    source = _make_resume(tmp_path / "source.docx")
    output = tmp_path / "unknown.docx"
    plan = _plan("bullet-does-not-exist", "Built a safe replacement sentence.")

    with pytest.raises(ResumeTailoringError, match="unknown bullet ID"):
        create_bullet_rewritten_docx(source, output, plan)

    assert not output.exists()


def test_cached_output_validation_rejects_non_target_text_change(tmp_path: Path) -> None:
    source = _make_resume(tmp_path / "source.docx")
    target = collect_editable_bullets(Document(source))[0]
    plan = _plan(
        target.bullet_id,
        "Built reliable AWS and Python pipelines for governed business reporting.",
    )
    output = create_bullet_rewritten_docx(source, tmp_path / "rewritten.docx", plan)
    tampered = Document(output)
    role = next(
        paragraph
        for paragraph in tampered.paragraphs
        if paragraph.text == "Senior Data Engineer | Example Company"
    )
    role.text = "Tampered role"
    tampered.save(output)

    with pytest.raises(ResumeTailoringError, match="Non-target resume content"):
        validate_bullet_rewritten_docx(source, output, plan)


def test_skips_complex_bullet_and_breaks_evidence_group(tmp_path: Path) -> None:
    badge = tmp_path / "badge.bmp"
    _write_bitmap(badge)
    complex_path = tmp_path / "complex.docx"
    complex_document = Document()
    complex_document.add_heading("Experience", level=1)
    complex_document.add_paragraph(
        "Built a safely editable platform before the unsupported paragraph.",
        style="List Bullet",
    )
    complex_bullet = complex_document.add_paragraph(style="List Bullet")
    complex_bullet.add_run("Built a platform with an inline diagram ")
    complex_bullet.add_run().add_picture(str(badge))
    complex_document.add_paragraph(
        "Automated a safely editable workflow after the unsupported paragraph.",
        style="List Bullet",
    )
    complex_document.save(complex_path)

    bullets = collect_editable_bullets(Document(complex_path))

    assert [bullet.bullet_id for bullet in bullets] == ["bullet-0001", "bullet-0003"]
    assert [bullet.group_id for bullet in bullets] == ["group-0001", "group-0002"]

    unsafe_plan = _plan(
        "bullet-0002",
        "Built a replacement that must not target the skipped complex paragraph.",
    )
    with pytest.raises(ResumeTailoringError, match="unknown bullet ID"):
        create_bullet_rewritten_docx(complex_path, tmp_path / "unsafe.docx", unsafe_plan)


def test_rejects_document_without_editable_bullets(tmp_path: Path) -> None:

    no_bullets_path = tmp_path / "no-bullets.docx"
    no_bullets = Document()
    no_bullets.add_heading("Summary", level=1)
    no_bullets.add_paragraph(
        "Built reliable platforms, but an action verb in the summary is not a bullet."
    )
    no_bullets.save(no_bullets_path)

    with pytest.raises(ResumeTailoringError, match="did not contain"):
        collect_editable_bullets(Document(no_bullets_path))
