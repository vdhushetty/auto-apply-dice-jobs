from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def resume_factory(tmp_path: Path):  # type: ignore[no-untyped-def]
    def create(name: str, profile: str, skills: list[str]) -> Path:
        path = tmp_path / name
        document = Document()
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(
            f"Data engineer with production experience delivering {profile} data platforms "
            "for analytics, machine learning, and reliable business reporting."
        )
        document.add_heading("Technical Skills", level=1)
        document.add_paragraph("Cloud: " + ", ".join(skills))
        document.add_heading("Experience", level=1)
        document.add_paragraph(
            "Built governed batch and streaming pipelines with monitored quality controls "
            "and repeatable infrastructure automation."
        )
        document.save(path)
        return path

    return create
