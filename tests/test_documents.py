from __future__ import annotations

import os
import struct
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement, parse_xml

from core.resumes.documents import (
    collect_skill_slots,
    create_curated_docx,
    fingerprint_document,
    validate_resume_path,
)
from core.resumes.models import ResumeConfigurationError

WORDPROCESSINGML_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
VML_NAMESPACE = "urn:schemas-microsoft-com:vml"


def _write_bitmap(path: Path, rgb: tuple[int, int, int]) -> None:
    red, green, blue = rgb
    path.write_bytes(
        struct.pack("<2sIHHI", b"BM", 58, 0, 0, 54)
        + struct.pack("<IiiHHIIiiII", 40, 1, 1, 1, 24, 0, 4, 2835, 2835, 0, 0)
        + bytes((blue, green, red, 0))
    )


def _add_legacy_drawing(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run._r.append(
        parse_xml(
            f'<w:pict xmlns:w="{WORDPROCESSINGML_NAMESPACE}" xmlns:v="{VML_NAMESPACE}">'
            '<v:shape id="synthetic-badge" style="width:1pt;height:1pt"/>'
            "</w:pict>"
        )
    )


def _write_docx_with_construct(path: Path, construct: str) -> None:
    document = Document()
    paragraph = document.add_paragraph("Synthetic resume content")
    if construct == "tracked_changes":
        paragraph._p.append(OxmlElement("w:ins"))
    elif construct == "content_controls":
        paragraph._p.append(OxmlElement("w:sdt"))
    elif construct == "text_boxes":
        paragraph.add_run()._r.append(
            parse_xml(
                f'<w:pict xmlns:w="{WORDPROCESSINGML_NAMESPACE}" xmlns:v="{VML_NAMESPACE}">'
                "<v:shape><v:textbox><w:txbxContent><w:p/></w:txbxContent></v:textbox>"
                "</v:shape></w:pict>"
            )
        )
    elif construct == "alt_chunk":
        paragraph._p.append(OxmlElement("w:altChunk"))
    else:
        raise AssertionError(f"Unknown test construct: {construct}")
    document.save(path)


def test_curated_docx_reorders_exact_items_and_preserves_structure(
    resume_factory,
    tmp_path: Path,  # type: ignore[no-untyped-def]
) -> None:
    source = resume_factory(
        "aws.docx",
        "AWS",
        ["AWS", "Python", "SQL", "Spark", "S3", "Glue"],
    )
    before_document = Document(source)
    before_fingerprint = fingerprint_document(before_document)
    slots = collect_skill_slots(before_document)
    assert len(slots) == 1
    slot = slots[0]
    reversed_order = tuple(reversed(slot.original_order))

    output = create_curated_docx(
        source,
        tmp_path / "tailored.docx",
        {slot.slot_id: reversed_order},
    )

    source_text = "\n".join(paragraph.text for paragraph in Document(source).paragraphs)
    output_document = Document(output)
    output_text = "\n".join(paragraph.text for paragraph in output_document.paragraphs)
    assert "Cloud: AWS, Python, SQL, Spark, S3, Glue" in source_text
    assert "Cloud: Glue, S3, Spark, SQL, Python, AWS" in output_text
    assert fingerprint_document(output_document) == before_fingerprint
    if os.name != "nt":
        assert output.stat().st_mode & 0o777 == 0o600
    assert sorted(item.text for item in slot.items) == sorted(
        ["AWS", "Python", "SQL", "Spark", "S3", "Glue"]
    )


def test_narrative_comma_list_is_not_editable(tmp_path: Path) -> None:
    source = tmp_path / "narrative.docx"
    document = Document()
    document.add_heading("Technical Skills", level=1)
    document.add_paragraph("Built Python pipelines, led SQL migration, optimized AWS costs")
    document.save(str(source))

    assert collect_skill_slots(Document(str(source))) == ()


def test_sample_shaped_skill_slot_preserves_terminal_period_and_objects(tmp_path: Path) -> None:
    source = tmp_path / "synthetic-sample.docx"
    badge = tmp_path / "synthetic-badge.bmp"
    _write_bitmap(badge, (20, 80, 160))

    document = Document()
    document.add_heading("Skills", level=1)
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run("Platforms:")
    label_run.bold = True
    body_run = paragraph.add_run(" AWS, Python, SQL, Spark.")
    body_run.bold = False
    document.add_picture(str(badge))
    _add_legacy_drawing(document)
    document.save(source)

    before_document = Document(source)
    slots = collect_skill_slots(before_document)
    assert len(slots) == 1
    slot = slots[0]
    assert slot.prefix == "Platforms: "
    assert slot.suffix == "."
    assert slot.body_run_indexes == (1,)
    assert [item.text for item in slot.items] == ["AWS", "Python", "SQL", "Spark"]

    before_fingerprint = fingerprint_document(before_document)
    assert len(before_fingerprint.embedded_media_hashes) == 1
    assert before_fingerprint.drawing_count == 1
    assert before_fingerprint.object_count == 1

    output = create_curated_docx(
        source,
        tmp_path / "synthetic-tailored.docx",
        {slot.slot_id: tuple(reversed(slot.original_order))},
    )

    output_document = Document(output)
    output_paragraph = output_document.paragraphs[1]
    assert output_paragraph.text == "Platforms: Spark, SQL, Python, AWS."
    assert output_paragraph.text.count(".") == 1
    assert output_paragraph.runs[0].bold
    assert output_paragraph.runs[1].bold is False
    assert fingerprint_document(output_document) == before_fingerprint


def test_non_taxonomy_list_inside_skills_section_is_editable(tmp_path: Path) -> None:
    source = tmp_path / "synthetic-reliability.docx"
    document = Document()
    document.add_heading("Skills", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Reliability & Delivery:").bold = True
    paragraph.add_run(
        " Monitoring, Data Quality, Incident Response, Release Management."
    ).bold = False
    document.save(source)

    slots = collect_skill_slots(Document(source))

    assert len(slots) == 1
    assert slots[0].suffix == "."
    assert len(slots[0].items) == 4


def test_fingerprint_hashes_embedded_media_contents(tmp_path: Path) -> None:
    fingerprints = []
    for name, color in (("blue", (0, 0, 255)), ("red", (255, 0, 0))):
        image_path = tmp_path / f"{name}.bmp"
        document_path = tmp_path / f"{name}.docx"
        _write_bitmap(image_path, color)
        document = Document()
        document.add_picture(str(image_path))
        document.save(document_path)
        fingerprints.append(fingerprint_document(Document(document_path)))

    assert fingerprints[0].drawing_count == fingerprints[1].drawing_count == 1
    assert len(fingerprints[0].embedded_media_hashes) == 1
    assert fingerprints[0].embedded_media_hashes != fingerprints[1].embedded_media_hashes


@pytest.mark.parametrize(
    ("construct", "message"),
    [
        ("tracked_changes", "tracked changes"),
        ("content_controls", "content controls"),
        ("text_boxes", "text boxes"),
        ("alt_chunk", "altChunk"),
    ],
)
def test_tailored_mode_rejects_unsupported_ooxml_constructs(
    tmp_path: Path,
    construct: str,
    message: str,
) -> None:
    source = tmp_path / f"{construct}.docx"
    _write_docx_with_construct(source, construct)

    assert validate_resume_path(source, tailored=False) == source.resolve()
    with pytest.raises(ResumeConfigurationError, match=message):
        validate_resume_path(source, tailored=True)
